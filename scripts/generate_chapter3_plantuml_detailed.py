from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PUML_DIR = DOCS / "plantuml"
DIAGRAM_DIR = DOCS / "diagrams"
PLANTUML_JAR = ROOT / "tools" / "plantuml.jar"
MAX_FIGURE_WIDTH_CM = 17.0
MAX_FIGURE_HEIGHT_CM = 21.0


def participants_for(kind: str) -> list[tuple[str, str, str]]:
    if kind == "ai":
        return [
            ("actor", "用户", "U"),
            ("participant", "AI 面板", "P"),
            ("participant", "前端 API", "A"),
            ("participant", "FastAPI 后端", "B"),
            ("participant", "豆包/火山方舟", "D"),
        ]
    if kind == "transcript":
        return [
            ("actor", "用户", "U"),
            ("participant", "成绩单导入组件", "I"),
            ("participant", "前端 API", "A"),
            ("participant", "FastAPI 后端", "B"),
            ("participant", "OCR/本地解析器", "O"),
            ("participant", "教育背景表单", "E"),
        ]
    if kind == "export":
        return [
            ("actor", "用户", "U"),
            ("participant", "Toolbar", "T"),
            ("participant", "简历预览 DOM", "R"),
            ("participant", "html2canvas", "C"),
            ("participant", "jsPDF", "P"),
            ("participant", "浏览器下载", "B"),
        ]
    if kind == "backend":
        return [
            ("actor", "前端工作台", "F"),
            ("participant", "FastAPI 应用", "B"),
            ("participant", "业务路由", "R"),
            ("participant", "业务服务", "S"),
            ("participant", "外部服务", "E"),
        ]
    return [
        ("actor", "用户", "U"),
        ("participant", "前端工作台", "F"),
        ("participant", "App 状态", "S"),
        ("participant", "localStorage", "L"),
        ("participant", "预览组件", "P"),
    ]


MODULES = [
    {
        "title": "3.1 用户工作台模块",
        "intro": "用户工作台模块负责材料创建、结构化录入、状态保存、主题与布局控制。该模块直接面向材料编写者，所有输入都会汇总到 ResumeData 数据结构，并驱动右侧预览与 AI 助手同步更新。",
        "items": [
            {
                "no": "3.1.1",
                "title": "进入系统与恢复工作台",
                "kind": "front",
                "desc": "用户打开系统首页，系统初始化默认简历数据、页面主题、模板、使用场景和 AI 面板状态。若浏览器本地存储中存在历史数据，系统读取并归一化后恢复到上次工作台状态；若读取失败，则使用默认数据进入系统。",
                "steps": [
                    ("U", "F", "打开系统页面"),
                    ("F", "L", "读取 resume-data、page-theme、template 等缓存"),
                    ("L", "F", "返回历史数据或空值"),
                    ("F", "S", "归一化 ResumeData 并设置当前状态"),
                    ("S", "P", "渲染编辑区与预览区"),
                ],
                "ok": "恢复工作台或显示默认入口",
                "fail": "缓存格式异常时丢弃异常数据并使用默认数据",
            },
            {
                "no": "3.1.2",
                "title": "选择材料使用场景",
                "kind": "front",
                "desc": "用户在总入口选择求职、考研、保研或升学场景，系统根据场景设置默认编辑标签、字段提示、模块顺序和目标输入名称，并打开工作台。场景信息会写入本地存储，方便刷新后继续编辑。",
                "steps": [
                    ("U", "F", "点击求职/考研/保研/升学入口"),
                    ("F", "S", "selectUseCase(useCase)"),
                    ("S", "S", "设置默认标签、场景提示和 AI 面板状态"),
                    ("S", "L", "保存 resume-use-case 与 workspace-started"),
                    ("S", "P", "按场景刷新预览标题和重点标签"),
                ],
                "ok": "进入对应材料工作台",
                "fail": "无效场景参数时保留默认场景",
            },
            {
                "no": "3.1.3",
                "title": "切换编辑模块",
                "kind": "front",
                "desc": "用户可在基本信息、教育背景、实践经历、项目科研、技能能力、奖项证书之间切换。系统会检查当前场景允许的 tabOrder，并保存活动标签；若缓存中的标签不属于当前场景，则切换到该场景默认标签。",
                "steps": [
                    ("U", "F", "点击编辑侧栏标签"),
                    ("F", "S", "updateActiveTab(tab)"),
                    ("S", "L", "保存 resume-active-tab"),
                    ("S", "F", "显示对应表单"),
                    ("S", "P", "保持预览同步"),
                ],
                "ok": "表单切换成功",
                "fail": "非法标签回退到默认编辑模块",
            },
            {
                "no": "3.1.4",
                "title": "基本信息维护",
                "kind": "front",
                "desc": "用户维护姓名、定位、邮箱、电话、所在地、籍贯、民族、性别、年龄和个人简介。系统在输入时更新 personalInfo 字段，并立即保存到本地。该功能不强制阻断空值，而是在评分与建议环节提示补充。",
                "steps": [
                    ("U", "F", "填写或修改基本信息字段"),
                    ("F", "S", "updatePersonalInfo(field, value)"),
                    ("S", "S", "克隆并更新 ResumeData.personalInfo"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新预览中的姓名、定位和简介"),
                ],
                "ok": "基本信息同步到材料预览",
                "fail": "空字段保留但后续评分提示补充",
            },
            {
                "no": "3.1.5",
                "title": "证件照上传与移除",
                "kind": "front",
                "desc": "用户选择证件照后，系统在浏览器端读取文件，按最大宽高压缩图片并转换为 base64 数据地址，写入 personalInfo.avatar。用户也可以点击移除按钮清空头像字段。",
                "steps": [
                    ("U", "F", "选择图片文件或点击移除"),
                    ("F", "F", "FileReader 读取图片"),
                    ("F", "F", "Canvas 等比压缩并生成 JPEG"),
                    ("F", "S", "更新 personalInfo.avatar"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新证件照展示"),
                ],
                "ok": "证件照显示在编辑区和模板中",
                "fail": "非图片或读取失败时不更新头像",
            },
            {
                "no": "3.1.6",
                "title": "教育背景维护",
                "kind": "front",
                "desc": "用户新增、编辑或删除教育经历，字段包括学校、专业、学历、起止时间、绩点、排名、英语水平、核心课程和说明。学术申请场景会突出绩点排名、英语水平和核心课程区域。",
                "steps": [
                    ("U", "F", "新增/修改/删除教育经历"),
                    ("F", "S", "updateEducation(index, field, value)"),
                    ("S", "S", "更新 education 列表"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新教育背景模块"),
                ],
                "ok": "教育经历进入简历预览",
                "fail": "删除最后一条后保留空教育列表并提示可新增",
            },
            {
                "no": "3.1.7",
                "title": "核心课程手动维护",
                "kind": "front",
                "desc": "用户可在教育经历下手动添加课程成绩，维护课程名称、成绩和学分。系统将课程保存到 education.coreCourses，用于保研、考研、升学模板展示和材料评分。",
                "steps": [
                    ("U", "F", "点击手动添加课程"),
                    ("F", "S", "addCourseScore(educationIndex)"),
                    ("U", "F", "填写课程、成绩、学分"),
                    ("F", "S", "updateCourseScore(...)"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新核心课程展示"),
                ],
                "ok": "课程成绩显示在教育背景中",
                "fail": "空课程不阻断保存，评分阶段提示精选高分课程",
            },
            {
                "no": "3.1.8",
                "title": "实践经历维护",
                "kind": "front",
                "desc": "用户维护公司、组织、实验室或社团经历，填写角色、起止时间和经历描述。系统支持多条经历的新增和删除，求职场景优先提示职责、工具和量化结果。",
                "steps": [
                    ("U", "F", "新增或编辑实践经历"),
                    ("F", "S", "updateExperience(index, field, value)"),
                    ("S", "S", "更新 experience 列表"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新经历模块"),
                ],
                "ok": "实践经历同步显示",
                "fail": "经历描述过短时由评分模块给出补充建议",
            },
            {
                "no": "3.1.9",
                "title": "项目科研维护",
                "kind": "front",
                "desc": "用户维护项目、科研、论文、课程作品等经历，填写项目名称、角色、起止时间、关键词和描述。系统将 technologies 字段转换为模板所需的 tags，保证多套模板兼容展示。",
                "steps": [
                    ("U", "F", "新增或编辑项目科研"),
                    ("F", "S", "updateProject(index, field, value)"),
                    ("S", "S", "维护 projects 与 technologies"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新项目科研模块"),
                ],
                "ok": "项目成果进入预览模板",
                "fail": "关键词为空时提示补充技术栈或研究方法",
            },
            {
                "no": "3.1.10",
                "title": "技能能力维护",
                "kind": "front",
                "desc": "用户新增技能、语言、工具或研究方法，并选择了解、熟练、精通、专家四级熟练度。技能列表会用于求职关键词匹配和材料评分。",
                "steps": [
                    ("U", "F", "添加技能并选择熟练度"),
                    ("F", "S", "updateSkill(index, field, value)"),
                    ("S", "S", "更新 skills 列表"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新技能展示"),
                ],
                "ok": "技能项进入材料预览",
                "fail": "技能少于建议数量时评分模块提醒补充",
            },
            {
                "no": "3.1.11",
                "title": "奖项证书维护",
                "kind": "front",
                "desc": "用户维护奖项、证书或荣誉，填写名称、颁发机构、获得时间和说明。系统保存 awards 列表，并在学术类场景优先展示奖学金、竞赛和证书信息。",
                "steps": [
                    ("U", "F", "新增或编辑奖项证书"),
                    ("F", "S", "updateAward(index, field, value)"),
                    ("S", "S", "更新 awards 列表"),
                    ("S", "L", "保存 resume-data"),
                    ("S", "P", "刷新奖项证书模块"),
                ],
                "ok": "奖项证书显示在材料中",
                "fail": "无奖项时不影响导出，但评分权重不加分",
            },
            {
                "no": "3.1.12",
                "title": "页面主题切换",
                "kind": "front",
                "desc": "用户可切换学术蓝、研究绿、论文紫三套页面主题。系统写入 page-theme，并将主题变量应用到工具栏、编辑区、预览区和 AI 面板。",
                "steps": [
                    ("U", "F", "选择页面主题"),
                    ("F", "S", "updateTheme(theme)"),
                    ("S", "L", "保存 page-theme"),
                    ("S", "F", "更新 CSS 变量"),
                    ("F", "P", "预览区同步主题背景"),
                ],
                "ok": "页面主题即时生效",
                "fail": "无效主题值回退到默认研究绿",
            },
            {
                "no": "3.1.13",
                "title": "工作区宽度调整",
                "kind": "front",
                "desc": "用户拖动编辑区和预览区之间的分隔条，系统根据容器宽度、AI 面板宽度和预览最小宽度计算编辑区宽度，并限制在安全范围内。",
                "steps": [
                    ("U", "F", "按下并拖动分隔条"),
                    ("F", "S", "startResize / resizePreview"),
                    ("S", "S", "计算 min/max editorWidth"),
                    ("S", "F", "更新 grid-template-columns"),
                    ("U", "F", "松开鼠标结束拖动"),
                ],
                "ok": "编辑区与预览区宽度稳定调整",
                "fail": "小屏设备隐藏分隔条并改为纵向布局",
            },
            {
                "no": "3.1.14",
                "title": "AI 助手展开与收起",
                "kind": "front",
                "desc": "用户点击 AI 面板侧边按钮后，系统切换 assistantOpen 状态，并保存到本地。展开状态显示完整 AI 助手，收起状态仅保留窄栏按钮，释放预览区域宽度。",
                "steps": [
                    ("U", "F", "点击 AI 助手开关"),
                    ("F", "S", "toggleAssistant()"),
                    ("S", "L", "保存 assistant-open"),
                    ("S", "F", "切换 assistant-shell 样式"),
                    ("F", "P", "重新计算工作台宽度"),
                ],
                "ok": "AI 面板展开或收起",
                "fail": "缓存失败时仍保持当前页面状态",
            },
            {
                "no": "3.1.15",
                "title": "本地保存与刷新恢复",
                "kind": "front",
                "desc": "用户每次更新材料内容，系统都会把完整 ResumeData 保存到 localStorage。页面刷新时，系统读取缓存并调用 normalizeResume 补齐缺失字段，保证旧版本数据也能兼容新版字段。",
                "steps": [
                    ("U", "F", "修改任意材料字段"),
                    ("F", "S", "updateResumeData(newData)"),
                    ("S", "L", "JSON.stringify 保存 resume-data"),
                    ("U", "F", "刷新页面"),
                    ("F", "L", "读取 resume-data"),
                    ("F", "S", "normalizeResume(data)"),
                ],
                "ok": "刷新后继续编辑上次材料",
                "fail": "JSON 解析失败时使用默认简历数据",
            },
        ],
    },
    {
        "title": "3.2 AI 辅助模块",
        "intro": "AI 辅助模块提供文本优化、简介生成、目标分析、材料匹配、职业建议和评分反馈。模块通过前端 AI 面板发起请求，FastAPI 后端封装云端模型调用，并提供本地兜底逻辑。",
        "items": [
            {
                "no": "3.2.1",
                "title": "选择优化字段",
                "kind": "ai",
                "desc": "用户在 AI 面板中选择个人简介、定位标题、实践经历或项目科研作为优化对象。系统根据当前 ResumeData 动态生成字段列表，若原字段被删除，则自动回退到可用字段。",
                "steps": [
                    ("U", "P", "打开智能优化标签页"),
                    ("P", "P", "根据 ResumeData 计算 optimizeFields"),
                    ("U", "P", "选择优化位置"),
                    ("P", "P", "校验字段是否存在且有内容"),
                    ("P", "U", "显示可优化字段"),
                ],
                "ok": "用户选中待优化内容",
                "fail": "内容为空时提示先填写材料",
            },
            {
                "no": "3.2.2",
                "title": "AI 文本优化",
                "kind": "ai",
                "desc": "用户点击优化文本，前端将内容和场景发送至 /api/optimize/。后端构造提示词调用豆包/火山方舟，要求保持原意、不编造经历，并突出动作、方法、结果和量化成果。",
                "steps": [
                    ("U", "P", "点击优化文本"),
                    ("P", "A", "optimizeSection(content, section)"),
                    ("A", "B", "POST /api/optimize/"),
                    ("B", "D", "调用模型生成优化文本"),
                    ("D", "B", "返回优化结果"),
                    ("B", "A", "success=true, optimized"),
                    ("A", "P", "展示结果和说明"),
                ],
                "ok": "用户可一键应用优化结果",
                "fail": "接口失败时返回原文并提示后端服务状态",
            },
            {
                "no": "3.2.3",
                "title": "个人简介生成",
                "kind": "ai",
                "desc": "用户点击生成简介后，前端将姓名、定位、原简介和经历摘要拼接为 resume_text，调用 /api/summary/。后端要求模型生成 50-120 字简介，并只返回可直接使用的文本。",
                "steps": [
                    ("U", "P", "点击生成简介"),
                    ("P", "A", "generateSummary(personalInfo, experience)"),
                    ("A", "B", "POST /api/summary/"),
                    ("B", "D", "请求生成专业简介"),
                    ("D", "B", "返回简介文本"),
                    ("B", "A", "summary"),
                    ("A", "P", "展示简介并等待应用"),
                ],
                "ok": "简介结果可写回 personalInfo.summary",
                "fail": "简历内容不足时返回提示或兜底摘要",
            },
            {
                "no": "3.2.4",
                "title": "目标要求分析",
                "kind": "ai",
                "desc": "用户粘贴岗位 JD、院校项目要求或导师方向后，点击分析目标。后端提取目标名称、行业/方向、必备技能、加分技能、职责和要求，帮助用户理解材料应重点突出哪些内容。",
                "steps": [
                    ("U", "P", "粘贴目标要求"),
                    ("P", "A", "analyzeJD(jd)"),
                    ("A", "B", "POST /api/jd/analyze"),
                    ("B", "D", "提取目标关键信息"),
                    ("D", "B", "返回结构化分析"),
                    ("B", "A", "result"),
                    ("A", "P", "展示目标分析"),
                ],
                "ok": "显示目标关键词和任职要求",
                "fail": "目标文本为空时提示先粘贴内容",
            },
            {
                "no": "3.2.5",
                "title": "简历与目标匹配",
                "kind": "ai",
                "desc": "用户点击匹配材料后，前端使用 resumeToText 把结构化简历转换为文本，与目标要求一起发送到 /api/jd/match。系统返回匹配分数、匹配关键词、缺失关键词和改进建议。",
                "steps": [
                    ("U", "P", "点击匹配材料"),
                    ("P", "A", "resumeToText(resume)"),
                    ("A", "B", "POST /api/jd/match"),
                    ("B", "D", "对比简历与目标要求"),
                    ("D", "B", "返回匹配分析"),
                    ("B", "A", "result"),
                    ("A", "P", "展示分数和建议"),
                ],
                "ok": "输出匹配报告",
                "fail": "模型失败时使用关键词匹配兜底",
            },
            {
                "no": "3.2.6",
                "title": "职业发展建议",
                "kind": "ai",
                "desc": "用户点击生成发展建议后，前端把当前简历转换为文本并调用 /api/career/。后端返回当前水平评估、短期目标、中期目标、推荐技能和发展路径建议。",
                "steps": [
                    ("U", "P", "点击生成发展建议"),
                    ("P", "A", "getCareerAdvice(resume)"),
                    ("A", "B", "POST /api/career/"),
                    ("B", "D", "生成职业发展路径"),
                    ("D", "B", "返回建议文本"),
                    ("B", "A", "result"),
                    ("A", "P", "展示发展建议"),
                ],
                "ok": "用户获得阶段性发展建议",
                "fail": "简历内容为空时提示先填写材料",
            },
            {
                "no": "3.2.7",
                "title": "材料完整度评分",
                "kind": "ai",
                "desc": "用户点击开始评分后，前端按规则统计基本信息、个人简介、教育、经历、项目、技能和奖项得分，并根据求职或学术场景追加针对性建议。该功能本地完成，不依赖后端。",
                "steps": [
                    ("U", "P", "点击开始评分"),
                    ("P", "P", "calculateResumeScore(resumeData)"),
                    ("P", "P", "按字段完整度累加分数"),
                    ("P", "P", "根据场景追加建议"),
                    ("P", "U", "展示分数、等级和建议"),
                ],
                "ok": "输出 0-100 分和改进建议",
                "fail": "材料过少时给出待完善等级",
            },
            {
                "no": "3.2.8",
                "title": "AI 调用失败兜底",
                "kind": "ai",
                "desc": "当云端 AI 出现鉴权、权限、额度、频率或网络问题时，后端根据错误类型返回中文提示，并使用本地规则生成兜底结果，保证用户仍能继续编辑材料。",
                "steps": [
                    ("P", "A", "发起 AI 请求"),
                    ("A", "B", "调用后端接口"),
                    ("B", "D", "请求云端模型"),
                    ("D", "B", "返回异常或超时"),
                    ("B", "B", "cloud_error_message + fallback"),
                    ("B", "A", "返回兜底结果和提示"),
                    ("A", "P", "展示可继续使用的结果"),
                ],
                "ok": "外部服务异常不阻断编辑流程",
                "fail": "严重异常时保留原文并提示检查配置",
            },
        ],
    },
    {
        "title": "3.3 成绩单识别模块",
        "intro": "成绩单识别模块用于从 PDF、docx 和图片中抽取课程成绩。系统优先使用后端结构化解析和阿里云 OCR；云端失败时，前端启用 PDF.js、mammoth 和 Tesseract.js 本地兜底。",
        "items": [
            {
                "no": "3.3.1",
                "title": "选择成绩单文件",
                "kind": "transcript",
                "desc": "用户在教育背景模块点击选择文件，可上传 PDF、doc、docx、png、jpg、jpeg、webp 等文件。组件初始化忙碌状态、清空旧候选课程，并提示正在读取成绩文件。",
                "steps": [
                    ("U", "I", "点击选择文件"),
                    ("I", "I", "打开隐藏 file input"),
                    ("U", "I", "上传成绩单文件"),
                    ("I", "I", "校验扩展名和 MIME 类型"),
                    ("I", "A", "准备 FormData"),
                ],
                "ok": "进入云端或本地解析流程",
                "fail": "不支持 .doc 旧格式时提示另存为 docx",
            },
            {
                "no": "3.3.2",
                "title": "Word 成绩单解析",
                "kind": "transcript",
                "desc": "当文件为 docx 时，后端读取 Word 压缩包中的 document.xml，解析表格行和段落文本，再按课程名、成绩、学分规则提取候选课程，避免把 Word 文档转成图片。",
                "steps": [
                    ("I", "A", "parseTranscriptFile(docx)"),
                    ("A", "B", "POST /api/transcript/parse"),
                    ("B", "O", "读取 document.xml"),
                    ("O", "B", "返回表格行与文本"),
                    ("B", "B", "parse_rows / parse_text_lines"),
                    ("B", "A", "courses/rawText/source=docx"),
                    ("A", "I", "显示候选课程"),
                ],
                "ok": "Word 成绩单解析完成",
                "fail": "文档结构异常时进入本地文本兜底",
            },
            {
                "no": "3.3.3",
                "title": "PDF 转图与云端识别",
                "kind": "transcript",
                "desc": "当文件为 PDF 时，后端使用 PyMuPDF 按页渲染为图片，压缩到云端 OCR 限制内，再调用阿里云统一识别的表格模式，抽取表格单元格和文本块。",
                "steps": [
                    ("I", "A", "parseTranscriptFile(pdf)"),
                    ("A", "B", "POST /api/transcript/parse"),
                    ("B", "B", "pdf_to_images + compress_image"),
                    ("B", "O", "调用阿里云表格 OCR"),
                    ("O", "B", "返回表格与文本"),
                    ("B", "B", "抽取课程并排序"),
                    ("B", "A", "返回候选课程"),
                ],
                "ok": "PDF 成绩单云端识别完成",
                "fail": "OCR 额度或权限异常时返回友好错误",
            },
            {
                "no": "3.3.4",
                "title": "图片成绩单云端识别",
                "kind": "transcript",
                "desc": "当用户上传图片成绩单时，后端先进行 EXIF 方向修正和尺寸压缩，再调用阿里云 OCR。系统从表格结构或普通文字中提取课程、成绩和学分。",
                "steps": [
                    ("I", "A", "parseTranscriptFile(image)"),
                    ("A", "B", "POST /api/transcript/parse"),
                    ("B", "B", "EXIF 转正与 JPEG 压缩"),
                    ("B", "O", "调用阿里云 OCR"),
                    ("O", "B", "返回识别结果"),
                    ("B", "B", "解析候选课程"),
                    ("B", "A", "返回 courses"),
                ],
                "ok": "图片成绩单识别完成",
                "fail": "图片不清晰或格式异常时提示更换文件",
            },
            {
                "no": "3.3.5",
                "title": "本地文本抽取兜底",
                "kind": "transcript",
                "desc": "当云端识别失败时，前端自动尝试本地兜底。PDF 使用 PDF.js 抽取文本，扫描页使用 Canvas 渲染后交给 Tesseract.js；docx 使用 mammoth 提取纯文本，再由本地规则解析课程成绩。",
                "steps": [
                    ("B", "A", "返回 success=false 或 courses 为空"),
                    ("A", "I", "提示云端失败并启动兜底"),
                    ("I", "O", "PDF.js/mammoth/Tesseract 解析文本"),
                    ("O", "I", "返回 rawText"),
                    ("I", "I", "parseCourseScores(text)"),
                    ("I", "U", "显示本地候选课程"),
                ],
                "ok": "云端失败后仍可获得候选课程",
                "fail": "本地也失败时显示两段错误信息",
            },
            {
                "no": "3.3.6",
                "title": "候选课程排序筛选",
                "kind": "transcript",
                "desc": "系统对识别到的课程按成绩优先、同分按学分、再按课程名排序。默认勾选前 6 门课程，用户可选择推荐 6 门、全选、清空或手动调整。",
                "steps": [
                    ("I", "I", "setCourseCandidates(courses)"),
                    ("I", "I", "sortCourses 按成绩和学分排序"),
                    ("I", "U", "展示课程、成绩、学分和勾选框"),
                    ("U", "I", "调整勾选状态"),
                    ("I", "I", "计算 selectedCount"),
                ],
                "ok": "用户确认要导入的课程集合",
                "fail": "未选中课程时禁用添加按钮",
            },
            {
                "no": "3.3.7",
                "title": "课程导入教育背景",
                "kind": "transcript",
                "desc": "用户点击添加选中课程后，组件向教育背景表单发出 addCourses 事件。系统过滤已存在的课程-成绩组合，只把新课程追加到当前教育经历的 coreCourses 中。",
                "steps": [
                    ("U", "I", "点击添加选中课程"),
                    ("I", "I", "生成 selected courses"),
                    ("I", "E", "emit addCourses(courses)"),
                    ("E", "E", "appendImportedCourses"),
                    ("E", "E", "过滤重复课程"),
                    ("E", "U", "提示已添加课程数量"),
                ],
                "ok": "课程成绩写入教育背景",
                "fail": "重复课程自动跳过",
            },
        ],
    },
    {
        "title": "3.4 预览导出与系统支撑模块",
        "intro": "该模块覆盖模板渲染、PDF 导出、后端接口聚合、环境变量配置和错误降级，是材料工作台稳定运行的支撑能力。",
        "items": [
            {
                "no": "3.4.1",
                "title": "模板选择",
                "kind": "front",
                "desc": "用户可选择现代通用、蓝金正式、深色侧栏、蓝灰、经典、现代、极简、时间线、专业等多套模板。系统保存 resume-template 并切换对应 Vue 模板组件。",
                "steps": [
                    ("U", "F", "在预览区选择模板"),
                    ("F", "S", "updateTemplate(template)"),
                    ("S", "L", "保存 resume-template"),
                    ("S", "P", "切换 currentTemplate"),
                    ("P", "U", "显示新模板效果"),
                ],
                "ok": "模板即时切换",
                "fail": "非法模板值回退到 modern",
            },
            {
                "no": "3.4.2",
                "title": "实时预览渲染",
                "kind": "front",
                "desc": "每次 ResumeData 更新后，预览器计算兼容字段，把 address、birthDate、politicalStatus、project.date、project.tags 等字段补齐，再传入当前模板组件生成 A4 预览。",
                "steps": [
                    ("S", "P", "传入 resumeData/theme/useCase/template"),
                    ("P", "P", "计算 compatibleResumeData"),
                    ("P", "P", "选择模板组件"),
                    ("P", "P", "渲染 .resume-preview"),
                    ("P", "U", "显示 A4 预览"),
                ],
                "ok": "编辑内容与预览保持同步",
                "fail": "缺失兼容字段时使用默认值补齐",
            },
            {
                "no": "3.4.3",
                "title": "PDF 导出",
                "kind": "export",
                "desc": "用户点击导出 PDF，系统查找 .resume-preview 节点，使用 html2canvas 按 2 倍缩放截图，再用 jsPDF 按画布尺寸生成文件并触发下载。",
                "steps": [
                    ("U", "T", "点击导出 PDF"),
                    ("T", "R", "查找 .resume-preview"),
                    ("T", "C", "html2canvas(resumeElement)"),
                    ("C", "T", "返回 canvas"),
                    ("T", "P", "addImage 并创建 PDF"),
                    ("P", "B", "save(场景标题.pdf)"),
                ],
                "ok": "浏览器下载 PDF 文件",
                "fail": "截图或写入失败时提示稍后重试",
            },
            {
                "no": "3.4.4",
                "title": "后端接口路由",
                "kind": "backend",
                "desc": "FastAPI 应用挂载 optimize、jd_analyzer、career、summary、transcript 五组路由，并启用 CORS 中间件。前端通过 VITE_API_BASE 指向后端服务地址。",
                "steps": [
                    ("F", "B", "发送 API 请求"),
                    ("B", "R", "按 prefix 匹配业务路由"),
                    ("R", "S", "调用业务服务函数"),
                    ("S", "E", "按需访问模型或 OCR"),
                    ("E", "S", "返回外部结果"),
                    ("S", "R", "组装响应模型"),
                    ("R", "F", "返回 JSON"),
                ],
                "ok": "前端获得结构化响应",
                "fail": "异常时返回 success=false 和错误说明",
            },
            {
                "no": "3.4.5",
                "title": "环境变量配置",
                "kind": "backend",
                "desc": "后端从 .env 读取 DOUBAO_API_KEY、DOUBAO_API_URL、DOUBAO_MODEL、阿里云 AccessKey 和 OCR endpoint。前端可通过 VITE_API_BASE 修改后端地址。",
                "steps": [
                    ("B", "B", "load_dotenv(.env)"),
                    ("B", "S", "读取豆包与阿里云配置"),
                    ("S", "S", "检查必要密钥"),
                    ("F", "B", "按 VITE_API_BASE 调用后端"),
                    ("B", "E", "携带配置访问外部服务"),
                ],
                "ok": "服务按配置正常运行",
                "fail": "缺少密钥时返回明确配置错误",
            },
            {
                "no": "3.4.6",
                "title": "错误提示与服务降级",
                "kind": "backend",
                "desc": "系统对前端请求失败、AI 服务不可用、OCR 权限不足、文件格式异常等情况提供中文提示。核心编辑、预览和导出能力不依赖云端服务，因此外部服务异常时仍能继续使用。",
                "steps": [
                    ("F", "B", "调用后端能力"),
                    ("B", "S", "捕获业务异常"),
                    ("S", "S", "转换为友好错误信息"),
                    ("S", "E", "外部服务失败时停止重试"),
                    ("S", "R", "返回兜底结果或错误消息"),
                    ("R", "F", "前端展示提示并保留已有数据"),
                ],
                "ok": "异常被用户可理解地呈现",
                "fail": "不可恢复错误不破坏本地编辑数据",
            },
        ],
    },
]


def ensure_dirs() -> None:
    DOCS.mkdir(exist_ok=True)
    if PUML_DIR.exists():
        shutil.rmtree(PUML_DIR)
    PUML_DIR.mkdir(parents=True)
    if DIAGRAM_DIR.exists():
        for file in DIAGRAM_DIR.glob("*.png"):
            file.unlink()
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)


OVERVIEW_FIGURES = [
    {
        "caption": "图 3-1 系统总体用例图",
        "stem": "figure_3_01_overall_use_case",
        "desc": "系统总体用例图从材料编写者视角描述系统边界内的主要业务能力，包括场景选择、结构化编辑、实时预览、AI 辅助、成绩单识别、模板切换、PDF 导出和本地持久化。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-1 系统总体用例图
actor "材料编写者" as User
rectangle "AI 简历与申请材料工作台" {
  usecase "选择材料场景" as UC1
  usecase "维护结构化简历" as UC2
  usecase "切换模板与主题" as UC3
  usecase "实时预览材料" as UC4
  usecase "调用 AI 辅助" as UC5
  usecase "导入成绩单课程" as UC6
  usecase "评估材料完整度" as UC7
  usecase "导出 PDF 材料" as UC8
  usecase "本地保存与恢复" as UC9
  usecase "服务异常降级" as UC10
}
User --> UC1
User --> UC2
User --> UC3
User --> UC4
User --> UC5
User --> UC6
User --> UC7
User --> UC8
UC2 .> UC9 : <<include>>
UC3 .> UC4 : <<include>>
UC5 .> UC10 : <<extend>>
UC6 .> UC10 : <<extend>>
UC8 .> UC4 : <<include>>
@enduml
""",
    },
    {
        "caption": "图 3-2 用户工作台模块用例图",
        "stem": "figure_3_02_workspace_use_case",
        "desc": "用户工作台模块用例图细化材料编辑入口，将场景、标签页、个人信息、教育经历、实践经历、项目科研、技能、奖项、证件照、布局宽度和 AI 面板状态纳入同一工作台边界。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 19
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 19
}
title 图 3-2 用户工作台模块用例图
actor "材料编写者" as User
rectangle "用户工作台模块" {
  usecase "恢复工作台" as W1
  usecase "选择使用场景" as W2
  usecase "切换编辑模块" as W3
  usecase "维护基本信息" as W4
  usecase "上传或移除证件照" as W5
  usecase "维护教育背景" as W6
  usecase "维护核心课程" as W7
  usecase "维护实践经历" as W8
  usecase "维护项目科研" as W9
  usecase "维护技能能力" as W10
  usecase "维护奖项证书" as W11
  usecase "切换页面主题" as W12
  usecase "调整工作区宽度" as W13
  usecase "展开或收起 AI 面板" as W14
  usecase "本地保存与刷新恢复" as W15
}
User --> W1
User --> W2
User --> W3
User --> W4
User --> W5
User --> W6
User --> W7
User --> W8
User --> W9
User --> W10
User --> W11
User --> W12
User --> W13
User --> W14
W4 .> W15 : <<include>>
W6 .> W15 : <<include>>
W8 .> W15 : <<include>>
W9 .> W15 : <<include>>
W10 .> W15 : <<include>>
W11 .> W15 : <<include>>
@enduml
""",
    },
    {
        "caption": "图 3-3 AI 辅助模块用例图",
        "stem": "figure_3_03_ai_use_case",
        "desc": "AI 辅助模块用例图描述智能优化、简介生成、目标分析、匹配分析、发展建议和材料评分等能力，同时体现云端模型失败后的本地兜底路径。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
title 图 3-3 AI 辅助模块用例图
actor "材料编写者" as User
actor "云端大模型服务" as LLM
rectangle "AI 辅助模块" {
  usecase "选择优化字段" as A1
  usecase "AI 文本优化" as A2
  usecase "生成个人简介" as A3
  usecase "分析目标要求" as A4
  usecase "匹配材料与目标" as A5
  usecase "生成发展建议" as A6
  usecase "材料完整度评分" as A7
  usecase "AI 调用失败兜底" as A8
}
User --> A1
User --> A2
User --> A3
User --> A4
User --> A5
User --> A6
User --> A7
A2 --> LLM
A3 --> LLM
A4 --> LLM
A5 --> LLM
A6 --> LLM
A2 .> A8 : <<extend>>
A3 .> A8 : <<extend>>
A4 .> A8 : <<extend>>
A5 .> A8 : <<extend>>
A6 .> A8 : <<extend>>
@enduml
""",
    },
    {
        "caption": "图 3-4 成绩单识别模块用例图",
        "stem": "figure_3_04_transcript_use_case",
        "desc": "成绩单识别模块用例图围绕文件选择、Word 解析、PDF 转图识别、图片 OCR、本地文本兜底、课程排序筛选和导入教育背景建立功能边界。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
title 图 3-4 成绩单识别模块用例图
actor "材料编写者" as User
actor "阿里云 OCR" as OCR
rectangle "成绩单识别模块" {
  usecase "选择成绩单文件" as T1
  usecase "Word 成绩单解析" as T2
  usecase "PDF 转图与云端识别" as T3
  usecase "图片成绩单云端识别" as T4
  usecase "本地文本抽取兜底" as T5
  usecase "候选课程排序筛选" as T6
  usecase "课程导入教育背景" as T7
}
User --> T1
User --> T6
User --> T7
T1 .> T2 : <<extend>>
T1 .> T3 : <<extend>>
T1 .> T4 : <<extend>>
T3 --> OCR
T4 --> OCR
T2 .> T5 : <<extend>>
T3 .> T5 : <<extend>>
T4 .> T5 : <<extend>>
T2 .> T6 : <<include>>
T3 .> T6 : <<include>>
T4 .> T6 : <<include>>
T5 .> T6 : <<include>>
T6 .> T7 : <<include>>
@enduml
""",
    },
    {
        "caption": "图 3-5 预览导出与系统支撑模块用例图",
        "stem": "figure_3_05_support_use_case",
        "desc": "预览导出与系统支撑模块用例图描述模板选择、实时预览、PDF 导出、后端接口路由、环境变量读取和错误降级之间的支撑关系。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
title 图 3-5 预览导出与系统支撑模块用例图
actor "材料编写者" as User
actor "浏览器下载器" as Browser
actor "外部服务" as Ext
rectangle "预览导出与系统支撑模块" {
  usecase "模板选择" as S1
  usecase "实时预览渲染" as S2
  usecase "PDF 导出" as S3
  usecase "后端接口路由" as S4
  usecase "环境变量配置" as S5
  usecase "错误提示与服务降级" as S6
}
User --> S1
User --> S2
User --> S3
S1 .> S2 : <<include>>
S3 .> S2 : <<include>>
S3 --> Browser
S4 --> Ext
S4 .> S5 : <<include>>
S4 .> S6 : <<extend>>
@enduml
""",
    },
    {
        "caption": "图 3-6 系统总体业务流程图",
        "stem": "figure_3_06_overall_flow",
        "desc": "系统总体业务流程图描述用户从进入工作台、恢复或创建材料，到编辑、预览、AI 辅助、成绩单导入、评分、导出和保存的完整业务链路。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam activity {
  BorderColor #334155
  BackgroundColor #F8FAFC
  DiamondBackgroundColor #E0F2FE
  DiamondBorderColor #0369A1
  ArrowColor #475569
  FontSize 20
}
title 图 3-6 系统总体业务流程图
start
:打开 AI 简历与申请材料工作台;
if (存在本地历史数据?) then (是)
  :读取 localStorage;
  :归一化 ResumeData;
else (否)
  :创建默认材料数据;
endif
:选择求职/考研/保研/升学场景;
:填写结构化材料内容;
:实时渲染 A4 预览;
if (需要 AI 辅助?) then (是)
  :选择优化、分析、匹配或建议功能;
  :调用后端 AI 接口;
  if (AI 服务可用?) then (是)
    :展示模型返回结果;
  else (否)
    :执行本地兜底并展示提示;
  endif
endif
if (需要导入成绩单?) then (是)
  :上传 PDF、Word 或图片成绩单;
  :识别候选课程并排序;
  :用户选择课程并导入教育背景;
endif
:进行材料完整度评分;
if (用户选择导出?) then (是)
  :生成 PDF 并触发浏览器下载;
else (否)
  :继续编辑;
endif
:保存最新工作台状态;
stop
@enduml
""",
    },
    {
        "caption": "图 3-7 AI 辅助处理流程图",
        "stem": "figure_3_07_ai_flow",
        "desc": "AI 辅助处理流程图描述前端校验、请求组装、后端提示词构造、云端模型调用、结果展示和异常兜底的处理顺序。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam activity {
  BorderColor #334155
  BackgroundColor #F8FAFC
  DiamondBackgroundColor #E0F2FE
  DiamondBorderColor #0369A1
  ArrowColor #475569
  FontSize 20
}
title 图 3-7 AI 辅助处理流程图
start
:用户在 AI 面板选择功能;
:读取当前 ResumeData 与场景;
if (输入内容满足要求?) then (是)
  :前端组装请求参数;
  :调用 FastAPI 对应接口;
  :后端构造中文提示词;
  if (外部模型配置有效?) then (是)
    :请求豆包/火山方舟模型;
    if (模型返回成功?) then (是)
      :解析模型文本或结构化结果;
    else (否)
      :生成本地兜底结果;
    endif
  else (否)
    :生成配置错误提示与兜底结果;
  endif
  :前端展示结果、说明和可应用按钮;
else (否)
  :提示用户先补充必要内容;
endif
stop
@enduml
""",
    },
    {
        "caption": "图 3-8 成绩单识别处理流程图",
        "stem": "figure_3_08_transcript_flow",
        "desc": "成绩单识别处理流程图描述不同文件类型下的解析分支，以及云端 OCR 失败后启用 PDF.js、mammoth 和 Tesseract.js 的本地兜底机制。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam activity {
  BorderColor #334155
  BackgroundColor #F8FAFC
  DiamondBackgroundColor #E0F2FE
  DiamondBorderColor #0369A1
  ArrowColor #475569
  FontSize 20
}
title 图 3-8 成绩单识别处理流程图
start
:用户选择成绩单文件;
:校验扩展名与 MIME 类型;
if (文件类型?) then (Word)
  :后端读取 document.xml;
  :解析表格行和段落文本;
elseif (PDF)
  :后端使用 PyMuPDF 转图片;
  :调用阿里云 OCR 表格识别;
else (图片)
  :修正 EXIF 方向并压缩图片;
  :调用阿里云 OCR;
endif
if (得到候选课程?) then (是)
  :按成绩、学分和课程名排序;
else (否)
  :前端启动本地解析兜底;
  :PDF.js/mammoth/Tesseract 提取文本;
  :本地规则解析课程成绩;
endif
:展示候选课程列表;
:默认勾选推荐课程;
if (用户确认导入?) then (是)
  :过滤重复课程;
  :追加到教育背景 coreCourses;
else (否)
  :保留候选结果等待调整;
endif
stop
@enduml
""",
    },
    {
        "caption": "图 3-9 PDF 导出处理流程图",
        "stem": "figure_3_09_pdf_export_flow",
        "desc": "PDF 导出处理流程图描述从预览节点检查、画布截图、PDF 页面生成到浏览器下载的过程，并包含预览节点缺失或生成失败时的异常提示。",
        "puml": """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam activity {
  BorderColor #334155
  BackgroundColor #F8FAFC
  DiamondBackgroundColor #E0F2FE
  DiamondBorderColor #0369A1
  ArrowColor #475569
  FontSize 20
}
title 图 3-9 PDF 导出处理流程图
start
:用户点击导出 PDF;
:查找 .resume-preview 节点;
if (预览节点存在?) then (是)
  :调用 html2canvas 生成高清画布;
  if (画布生成成功?) then (是)
    :按画布尺寸创建 jsPDF 文档;
    :写入预览截图;
    :按场景标题命名文件;
    :触发浏览器下载;
  else (否)
    :提示截图失败并保留当前页面;
  endif
else (否)
  :提示未找到可导出的预览内容;
endif
stop
@enduml
""",
    },
]


def apply_compact_use_case_templates() -> None:
    OVERVIEW_FIGURES[0]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype ortho
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 18
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 18
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 18
}
title 图 3-1 系统总体用例图
actor "材料编写者" as User
rectangle "AI 简历与申请材料工作台" {
  package "用户工作台" {
    usecase "管理材料工作台" as O1
    usecase "选择材料场景" as O11
    usecase "维护结构化简历" as O12
    usecase "本地保存与恢复" as O13
  }
  package "AI 辅助" {
    usecase "调用 AI 辅助服务" as O2
    usecase "文本优化/简介生成" as O21
    usecase "目标分析/材料匹配" as O22
    usecase "完整度评分" as O23
  }
  package "成绩单识别" {
    usecase "导入成绩单课程" as O3
    usecase "上传成绩单文件" as O31
    usecase "识别候选课程" as O32
    usecase "导入教育背景" as O33
  }
  package "预览导出与支撑" {
    usecase "生成申请材料" as O4
    usecase "切换模板与主题" as O41
    usecase "实时预览材料" as O42
    usecase "导出 PDF 材料" as O43
    usecase "服务异常降级" as O44
  }
}
User --> O1
User --> O2
User --> O3
User --> O4
O1 .> O11 : <<include>>
O1 .> O12 : <<include>>
O1 .> O13 : <<include>>
O2 .> O21 : <<include>>
O2 .> O22 : <<include>>
O2 .> O23 : <<include>>
O2 .> O44 : <<extend>>
O3 .> O31 : <<include>>
O3 .> O32 : <<include>>
O3 .> O33 : <<include>>
O3 .> O44 : <<extend>>
O4 .> O41 : <<include>>
O4 .> O42 : <<include>>
O4 .> O43 : <<include>>
O43 .> O42 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[1]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype ortho
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 17
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 17
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 17
}
title 图 3-2 用户工作台模块用例图
actor "材料编写者" as User
rectangle "用户工作台模块" {
  usecase "维护材料工作台" as W0
  package "资料内容维护" {
    usecase "维护基本信息" as W1
    usecase "上传或移除证件照" as W2
    usecase "维护教育背景" as W3
    usecase "维护核心课程" as W4
    usecase "维护实践经历" as W5
    usecase "维护项目科研" as W6
    usecase "维护技能能力" as W7
    usecase "维护奖项证书" as W8
  }
  package "工作台控制" {
    usecase "恢复工作台" as W9
    usecase "选择使用场景" as W10
    usecase "切换编辑模块" as W11
    usecase "切换页面主题" as W12
    usecase "调整工作区宽度" as W13
    usecase "展开或收起 AI 面板" as W14
  }
  usecase "本地保存与刷新恢复" as W15
}
User --> W0
W0 .> W1 : <<include>>
W0 .> W2 : <<include>>
W0 .> W3 : <<include>>
W0 .> W4 : <<include>>
W0 .> W5 : <<include>>
W0 .> W6 : <<include>>
W0 .> W7 : <<include>>
W0 .> W8 : <<include>>
User --> W9
User --> W10
User --> W11
User --> W12
User --> W13
User --> W14
W0 .> W15 : <<include>>
W9 .> W15 : <<include>>
W10 .> W15 : <<include>>
W11 .> W15 : <<include>>
W12 .> W15 : <<include>>
W14 .> W15 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[2]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype ortho
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 18
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 18
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 18
}
title 图 3-3 AI 辅助模块用例图
actor "材料编写者" as User
actor "云端大模型服务" as LLM
rectangle "AI 辅助模块" {
  usecase "使用 AI 助手" as A0
  package "内容生成与优化" {
    usecase "选择优化字段" as A1
    usecase "AI 文本优化" as A2
    usecase "生成个人简介" as A3
  }
  package "目标理解与匹配" {
    usecase "分析目标要求" as A4
    usecase "匹配材料与目标" as A5
    usecase "生成发展建议" as A6
  }
  package "本地反馈" {
    usecase "材料完整度评分" as A7
    usecase "AI 调用失败兜底" as A8
  }
}
User --> A0
A0 .> A1 : <<include>>
A0 .> A2 : <<include>>
A0 .> A3 : <<include>>
A0 .> A4 : <<include>>
A0 .> A5 : <<include>>
A0 .> A6 : <<include>>
A0 .> A7 : <<include>>
A2 --> LLM
A3 --> LLM
A4 --> LLM
A5 --> LLM
A6 --> LLM
A2 .> A8 : <<extend>>
A3 .> A8 : <<extend>>
A4 .> A8 : <<extend>>
A5 .> A8 : <<extend>>
A6 .> A8 : <<extend>>
@enduml
"""
    OVERVIEW_FIGURES[3]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype ortho
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 18
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 18
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 18
}
title 图 3-4 成绩单识别模块用例图
actor "材料编写者" as User
actor "阿里云 OCR" as OCR
rectangle "成绩单识别模块" {
  usecase "导入成绩单课程" as T0
  package "文件解析" {
    usecase "选择成绩单文件" as T1
    usecase "Word 成绩单解析" as T2
    usecase "PDF 转图与云端识别" as T3
    usecase "图片成绩单云端识别" as T4
  }
  package "结果处理" {
    usecase "本地文本抽取兜底" as T5
    usecase "候选课程排序筛选" as T6
    usecase "课程导入教育背景" as T7
  }
}
User --> T0
T0 .> T1 : <<include>>
T1 .> T2 : <<extend>>
T1 .> T3 : <<extend>>
T1 .> T4 : <<extend>>
T3 --> OCR
T4 --> OCR
T2 .> T5 : <<extend>>
T3 .> T5 : <<extend>>
T4 .> T5 : <<extend>>
T0 .> T6 : <<include>>
T0 .> T7 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[4]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype ortho
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 18
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 18
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 18
}
title 图 3-5 预览导出与系统支撑模块用例图
actor "材料编写者" as User
actor "浏览器下载器" as Browser
actor "外部服务" as Ext
rectangle "预览导出与系统支撑模块" {
  usecase "生成可交付材料" as S0
  package "前端呈现" {
    usecase "模板选择" as S1
    usecase "实时预览渲染" as S2
    usecase "PDF 导出" as S3
  }
  package "后端支撑" {
    usecase "后端接口路由" as S4
    usecase "环境变量配置" as S5
    usecase "错误提示与服务降级" as S6
  }
}
User --> S0
S0 .> S1 : <<include>>
S0 .> S2 : <<include>>
S0 .> S3 : <<include>>
S3 --> Browser
S4 --> Ext
S4 .> S5 : <<include>>
S4 .> S6 : <<extend>>
S0 .> S4 : <<include>>
S0 .> S6 : <<extend>>
@enduml
"""


def apply_report_use_case_templates() -> None:
    OVERVIEW_FIGURES[0]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-1 系统总体用例图
actor "材料编写者" as User
actor "云端 AI 服务" as AI
actor "OCR 服务" as OCR
actor "浏览器下载器" as Browser
rectangle "AI 简历与申请材料工作台" {
  usecase "选择材料场景\\n求职/考研/保研/升学" as UC1
  usecase "维护结构化材料\\n个人信息/教育/经历/项目/技能/奖项" as UC2
  usecase "实时预览与模板控制\\n模板/主题/宽度/AI 面板" as UC3
  usecase "AI 辅助处理\\n优化/简介/分析/匹配/建议/评分" as UC4
  usecase "成绩单识别导入\\nWord/PDF/图片/本地兜底" as UC5
  usecase "导出 PDF 材料\\n截图生成/文件下载" as UC6
  usecase "本地保存与刷新恢复\\nlocalStorage/数据归一化" as UC7
  usecase "服务异常降级\\n友好提示/兜底结果" as UC8
}
User --> UC1
User --> UC2
User --> UC3
User --> UC4
User --> UC5
User --> UC6
User --> UC7
AI --> UC4
OCR --> UC5
Browser --> UC6
UC4 .> UC8 : <<extend>>
UC5 .> UC8 : <<extend>>
UC2 .> UC7 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[1]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-2 用户工作台模块用例图
actor "材料编写者" as User
rectangle "用户工作台模块" {
  usecase "进入与恢复工作台\\n初始化数据/恢复历史/选择场景" as W1
  usecase "维护个人资料\\n基本信息/证件照" as W2
  usecase "维护教育信息\\n教育背景/核心课程" as W3
  usecase "维护经历成果\\n实践经历/项目科研/奖项证书" as W4
  usecase "维护能力标签\\n技能能力/熟练度" as W5
  usecase "控制工作台界面\\n编辑模块/主题/宽度/AI 面板" as W6
  usecase "本地保存与刷新恢复\\nresume-data/page-theme/template" as W7
}
User --> W1
User --> W2
User --> W3
User --> W4
User --> W5
User --> W6
W2 .> W7 : <<include>>
W3 .> W7 : <<include>>
W4 .> W7 : <<include>>
W5 .> W7 : <<include>>
W6 .> W7 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[2]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-3 AI 辅助模块用例图
actor "材料编写者" as User
actor "云端大模型服务" as LLM
rectangle "AI 辅助模块" {
  usecase "选择优化字段\\n简介/定位/经历/项目" as A1
  usecase "AI 文本优化\\n动作/方法/结果/量化成果" as A2
  usecase "生成个人简介\\n50-120 字专业总结" as A3
  usecase "分析目标要求\\n岗位 JD/院校项目/导师方向" as A4
  usecase "匹配材料与目标\\n分数/关键词/缺口/建议" as A5
  usecase "生成发展建议\\n短期目标/中期目标/技能路径" as A6
  usecase "材料完整度评分\\n本地规则/场景建议" as A7
  usecase "AI 调用失败兜底\\n中文错误提示/本地结果" as A8
}
User --> A1
User --> A2
User --> A3
User --> A4
User --> A5
User --> A6
User --> A7
LLM --> A2
LLM --> A3
LLM --> A4
LLM --> A5
LLM --> A6
A2 .> A8 : <<extend>>
A3 .> A8 : <<extend>>
A4 .> A8 : <<extend>>
A5 .> A8 : <<extend>>
A6 .> A8 : <<extend>>
@enduml
"""
    OVERVIEW_FIGURES[3]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-4 成绩单识别模块用例图
actor "材料编写者" as User
actor "阿里云 OCR" as OCR
rectangle "成绩单识别模块" {
  usecase "选择成绩单文件\\nPDF/docx/图片格式校验" as T1
  usecase "Word 成绩单解析\\ndocument.xml/表格/段落" as T2
  usecase "PDF 转图与云端识别\\nPyMuPDF/表格 OCR" as T3
  usecase "图片成绩单识别\\nEXIF 修正/压缩/OCR" as T4
  usecase "本地文本抽取兜底\\nPDF.js/mammoth/Tesseract" as T5
  usecase "候选课程排序筛选\\n成绩/学分/课程名/勾选" as T6
  usecase "课程导入教育背景\\n去重/追加 coreCourses" as T7
}
User --> T1
User --> T6
User --> T7
OCR --> T3
OCR --> T4
T1 .> T2 : <<extend>>
T1 .> T3 : <<extend>>
T1 .> T4 : <<extend>>
T2 .> T5 : <<extend>>
T3 .> T5 : <<extend>>
T4 .> T5 : <<extend>>
T2 .> T6 : <<include>>
T3 .> T6 : <<include>>
T4 .> T6 : <<include>>
T5 .> T6 : <<include>>
T6 .> T7 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[4]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-5 预览导出与系统支撑模块用例图
actor "材料编写者" as User
actor "浏览器下载器" as Browser
actor "外部服务" as Ext
rectangle "预览导出与系统支撑模块" {
  usecase "模板选择\\n多模板切换/模板缓存" as S1
  usecase "实时预览渲染\\n字段兼容/A4 预览" as S2
  usecase "PDF 导出\\nhtml2canvas/jsPDF/下载" as S3
  usecase "后端接口路由\\noptimize/jd/career/summary/transcript" as S4
  usecase "环境变量配置\\.env/VITE_API_BASE/密钥检查" as S5
  usecase "错误提示与服务降级\\n中文提示/保留本地数据" as S6
}
User --> S1
User --> S2
User --> S3
Browser --> S3
Ext --> S4
S1 .> S2 : <<include>>
S3 .> S2 : <<include>>
S4 .> S5 : <<include>>
S4 .> S6 : <<extend>>
@enduml
"""


def apply_report_flow_templates() -> None:
    OVERVIEW_FIGURES[5]["desc"] = "系统总体业务流程图描述用户从进入工作台、选择场景、编辑材料、调用辅助能力，到预览检查、导出 PDF 和保存状态的主流程；AI、成绩单识别和 PDF 导出的内部细节在后续流程图中展开。"
    OVERVIEW_FIGURES[5]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
skinparam dpi 180
skinparam shadowing false
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 22
skinparam titleFontSize 28
skinparam activity {
  BorderColor #334155
  BackgroundColor #F8FAFC
  DiamondBackgroundColor #E0F2FE
  DiamondBorderColor #0369A1
  ArrowColor #475569
  FontSize 22
}
title 图 3-6 系统总体业务流程图
start
:打开工作台并初始化数据;
if (存在本地历史数据?) then (是)
  :恢复上次材料与界面状态;
else (否)
  :创建默认材料数据;
endif
:选择求职/考研/保研/升学场景;
:维护结构化材料内容;
:实时预览并调整模板主题;
if (需要辅助处理?) then (是)
  :调用 AI 辅助、成绩单识别或评分功能;
else (否)
  :直接检查当前材料效果;
endif
if (确认导出 PDF?) then (是)
  :生成 PDF 并触发浏览器下载;
else (否)
  :继续编辑当前材料;
endif
:保存最新工作台状态;
stop
@enduml
"""


def apply_ordered_use_case_templates() -> None:
    OVERVIEW_FIGURES[0]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype polyline
skinparam nodesep 70
skinparam ranksep 90
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-1 系统总体用例图
actor "材料编写者" as User
rectangle "AI 简历与申请材料工作台" {
  usecase "选择材料场景\\n求职/考研/保研/升学" as UC1
  usecase "维护结构化材料\\n个人信息/教育/经历/项目/技能/奖项" as UC2
  usecase "实时预览与模板控制\\n模板/主题/宽度/AI 面板" as UC3
  usecase "AI 辅助处理\\n优化/简介/分析/匹配/建议/评分" as UC4
  usecase "成绩单识别导入\\nWord/PDF/图片/本地兜底" as UC5
  usecase "导出 PDF 材料\\n截图生成/文件下载" as UC6
  usecase "本地保存与刷新恢复\\nlocalStorage/数据归一化" as UC7
  usecase "服务异常降级\\n友好提示/兜底结果" as UC8
  UC1 -[hidden]down- UC2
  UC2 -[hidden]down- UC3
  UC3 -[hidden]down- UC4
  UC4 -[hidden]down- UC5
  UC5 -[hidden]down- UC6
  UC6 -[hidden]down- UC7
  UC7 -[hidden]down- UC8
}
actor "云端 AI 服务" as AI
actor "OCR 服务" as OCR
actor "浏览器下载器" as Browser
User -- UC1
User -- UC2
User -- UC3
User -- UC4
User -- UC5
User -- UC6
User -- UC7
UC4 -- AI
UC5 -- OCR
UC6 -- Browser
UC2 .down.> UC7 : <<include>>
UC4 .down.> UC8 : <<extend>>
UC5 .down.> UC8 : <<extend>>
@enduml
"""
    OVERVIEW_FIGURES[1]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype polyline
skinparam nodesep 70
skinparam ranksep 90
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-2 用户工作台模块用例图
actor "材料编写者" as User
rectangle "用户工作台模块" {
  usecase "进入与恢复工作台\\n初始化数据/恢复历史/选择场景" as W1
  usecase "维护个人资料\\n基本信息/证件照" as W2
  usecase "维护教育信息\\n教育背景/核心课程" as W3
  usecase "维护经历成果\\n实践经历/项目科研/奖项证书" as W4
  usecase "维护能力标签\\n技能能力/熟练度" as W5
  usecase "控制工作台界面\\n编辑模块/主题/宽度/AI 面板" as W6
  usecase "本地保存与刷新恢复\\nresume-data/page-theme/template" as W7
  W1 -[hidden]down- W2
  W2 -[hidden]down- W3
  W3 -[hidden]down- W4
  W4 -[hidden]down- W5
  W5 -[hidden]down- W6
  W6 -[hidden]down- W7
}
User -- W1
User -- W2
User -- W3
User -- W4
User -- W5
User -- W6
User -- W7
W1 .down.> W7 : <<include>>
W2 .down.> W7 : <<include>>
W3 .down.> W7 : <<include>>
W4 .down.> W7 : <<include>>
W5 .down.> W7 : <<include>>
W6 .down.> W7 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[2]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype polyline
skinparam nodesep 70
skinparam ranksep 90
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-3 AI 辅助模块用例图
actor "材料编写者" as User
rectangle "AI 辅助模块" {
  usecase "选择优化字段\\n简介/定位/经历/项目" as A1
  usecase "AI 文本优化\\n动作/方法/结果/量化成果" as A2
  usecase "生成个人简介\\n50-120 字专业总结" as A3
  usecase "分析目标要求\\n岗位 JD/院校项目/导师方向" as A4
  usecase "匹配材料与目标\\n分数/关键词/缺口/建议" as A5
  usecase "生成发展建议\\n短期目标/中期目标/技能路径" as A6
  usecase "材料完整度评分\\n本地规则/场景建议" as A7
  usecase "AI 调用失败兜底\\n中文错误提示/本地结果" as A8
  A1 -[hidden]down- A2
  A2 -[hidden]down- A3
  A3 -[hidden]down- A4
  A4 -[hidden]down- A5
  A5 -[hidden]down- A6
  A6 -[hidden]down- A7
  A7 -[hidden]down- A8
}
actor "云端大模型服务" as LLM
User -- A1
User -- A2
User -- A3
User -- A4
User -- A5
User -- A6
User -- A7
A2 -- LLM
A3 -- LLM
A4 -- LLM
A5 -- LLM
A6 -- LLM
A2 .down.> A8 : <<extend>>
A3 .down.> A8 : <<extend>>
A4 .down.> A8 : <<extend>>
A5 .down.> A8 : <<extend>>
A6 .down.> A8 : <<extend>>
@enduml
"""
    OVERVIEW_FIGURES[3]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype polyline
skinparam nodesep 70
skinparam ranksep 90
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-4 成绩单识别模块用例图
actor "材料编写者" as User
rectangle "成绩单识别模块" {
  usecase "选择成绩单文件\\nPDF/docx/图片格式校验" as T1
  usecase "Word 成绩单解析\\ndocument.xml/表格/段落" as T2
  usecase "PDF 转图与云端识别\\nPyMuPDF/表格 OCR" as T3
  usecase "图片成绩单识别\\nEXIF 修正/压缩/OCR" as T4
  usecase "本地文本抽取兜底\\nPDF.js/mammoth/Tesseract" as T5
  usecase "候选课程排序筛选\\n成绩/学分/课程名/勾选" as T6
  usecase "课程导入教育背景\\n去重/追加 coreCourses" as T7
  T1 -[hidden]down- T2
  T2 -[hidden]down- T3
  T3 -[hidden]down- T4
  T4 -[hidden]down- T5
  T5 -[hidden]down- T6
  T6 -[hidden]down- T7
}
actor "阿里云 OCR" as OCR
User -- T1
User -- T6
User -- T7
T3 -- OCR
T4 -- OCR
T1 .down.> T2 : <<extend>>
T1 .down.> T3 : <<extend>>
T1 .down.> T4 : <<extend>>
T2 .down.> T5 : <<extend>>
T3 .down.> T5 : <<extend>>
T4 .down.> T5 : <<extend>>
T2 .down.> T6 : <<include>>
T3 .down.> T6 : <<include>>
T4 .down.> T6 : <<include>>
T5 .down.> T6 : <<include>>
T6 .down.> T7 : <<include>>
@enduml
"""
    OVERVIEW_FIGURES[4]["puml"] = """@startuml
' Generated by scripts/generate_chapter3_plantuml_detailed.py
left to right direction
skinparam dpi 180
skinparam shadowing false
skinparam linetype polyline
skinparam nodesep 70
skinparam ranksep 90
skinparam defaultFontName Microsoft YaHei
skinparam defaultFontSize 20
skinparam titleFontSize 26
skinparam packageStyle rectangle
skinparam usecase {
  BorderColor #334155
  BackgroundColor #F8FAFC
  ArrowColor #475569
  FontSize 20
}
skinparam actor {
  BorderColor #334155
  BackgroundColor #E0F2FE
  FontSize 20
}
title 图 3-5 预览导出与系统支撑模块用例图
actor "材料编写者" as User
rectangle "预览导出与系统支撑模块" {
  usecase "模板选择\\n多模板切换/模板缓存" as S1
  usecase "实时预览渲染\\n字段兼容/A4 预览" as S2
  usecase "PDF 导出\\nhtml2canvas/jsPDF/下载" as S3
  usecase "后端接口路由\\noptimize/jd/career/summary/transcript" as S4
  usecase "环境变量配置\\.env/VITE_API_BASE/密钥检查" as S5
  usecase "错误提示与服务降级\\n中文提示/保留本地数据" as S6
  S1 -[hidden]down- S2
  S2 -[hidden]down- S3
  S3 -[hidden]down- S4
  S4 -[hidden]down- S5
  S5 -[hidden]down- S6
}
actor "浏览器下载器" as Browser
actor "外部服务" as Ext
User -- S1
User -- S2
User -- S3
S3 -- Browser
S4 -- Ext
S1 .down.> S2 : <<include>>
S3 .up.> S2 : <<include>>
S4 .down.> S5 : <<include>>
S4 .down.> S6 : <<extend>>
@enduml
"""


def apply_dot_use_case_templates() -> None:
    actor_user = 'label=<<FONT POINT-SIZE="26">&#9675;</FONT><BR/><FONT POINT-SIZE="26">&#47;|&#92;</FONT><BR/><FONT POINT-SIZE="26">&#47; &#92;</FONT><BR/>材料编写者>'
    actor_ai = 'label=<<FONT POINT-SIZE="26">&#9675;</FONT><BR/><FONT POINT-SIZE="26">&#47;|&#92;</FONT><BR/><FONT POINT-SIZE="26">&#47; &#92;</FONT><BR/>云端 AI 服务>'
    actor_llm = 'label=<<FONT POINT-SIZE="26">&#9675;</FONT><BR/><FONT POINT-SIZE="26">&#47;|&#92;</FONT><BR/><FONT POINT-SIZE="26">&#47; &#92;</FONT><BR/>云端大模型服务>'
    actor_ocr = 'label=<<FONT POINT-SIZE="26">&#9675;</FONT><BR/><FONT POINT-SIZE="26">&#47;|&#92;</FONT><BR/><FONT POINT-SIZE="26">&#47; &#92;</FONT><BR/>OCR 服务>'
    actor_browser = 'label=<<FONT POINT-SIZE="26">&#9675;</FONT><BR/><FONT POINT-SIZE="26">&#47;|&#92;</FONT><BR/><FONT POINT-SIZE="26">&#47; &#92;</FONT><BR/>浏览器下载器>'
    actor_ext = 'label=<<FONT POINT-SIZE="26">&#9675;</FONT><BR/><FONT POINT-SIZE="26">&#47;|&#92;</FONT><BR/><FONT POINT-SIZE="26">&#47; &#92;</FONT><BR/>外部服务>'

    OVERVIEW_FIGURES[0]["puml"] = f"""@startdot
digraph G {{
  graph [
    label="图 3-1 系统总体用例图",
    labelloc=t,
    rankdir=LR,
    splines=ortho,
    nodesep=0.45,
    ranksep=1.0,
    pad=0.25,
    fontname="Microsoft YaHei",
    fontsize=28
  ];
  node [fontname="Microsoft YaHei", fontsize=20, color="#334155", penwidth=1.6];
  edge [fontname="Microsoft YaHei", fontsize=18, color="#475569", penwidth=1.8, arrowsize=0.8];

  User [shape=plain, {actor_user}];
  AI [shape=plain, {actor_ai}];
  OCR [shape=plain, {actor_ocr}];
  Browser [shape=plain, {actor_browser}];

  subgraph cluster_system {{
    label="AI 简历与申请材料工作台";
    color="#334155";
    penwidth=1.8;
    fontsize=24;
    style="rounded";
    UC1 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="选择材料场景\\n求职/考研/保研/升学"];
    UC2 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="维护结构化材料\\n个人信息/教育/经历/项目/技能/奖项"];
    UC3 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="实时预览与模板控制\\n模板/主题/宽度/AI 面板"];
    UC4 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="AI 辅助处理\\n优化/简介/分析/匹配/建议/评分"];
    UC5 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="成绩单识别导入\\nWord/PDF/图片/本地兜底"];
    UC6 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="导出 PDF 材料\\n截图生成/文件下载"];
    UC7 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="本地保存与刷新恢复\\nlocalStorage/数据归一化"];
    UC8 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="服务异常降级\\n友好提示/兜底结果"];
    {{ rank=same; UC1; UC2; UC3; UC4; UC5; UC6; UC7; UC8; }}
    UC1 -> UC2 [style=invis, weight=50];
    UC2 -> UC3 [style=invis, weight=50];
    UC3 -> UC4 [style=invis, weight=50];
    UC4 -> UC5 [style=invis, weight=50];
    UC5 -> UC6 [style=invis, weight=50];
    UC6 -> UC7 [style=invis, weight=50];
    UC7 -> UC8 [style=invis, weight=50];
  }}

  {{ rank=source; User; }}
  {{ rank=sink; AI; OCR; Browser; }}
  User -> UC1 [dir=none];
  User -> UC2 [dir=none];
  User -> UC3 [dir=none];
  User -> UC4 [dir=none];
  User -> UC5 [dir=none];
  User -> UC6 [dir=none];
  User -> UC7 [dir=none];
  UC4 -> AI [dir=none];
  UC5 -> OCR [dir=none];
  UC6 -> Browser [dir=none];
  UC2 -> UC7 [style=dashed, arrowhead=open, label="<<include>>"];
  UC3 -> UC7 [style=dashed, arrowhead=open, label="<<include>>"];
  UC6 -> UC3 [style=dashed, arrowhead=open, label="<<include>>"];
  UC4 -> UC8 [style=dashed, arrowhead=open, label="<<extend>>"];
  UC5 -> UC8 [style=dashed, arrowhead=open, label="<<extend>>"];
}}
@enddot
"""
    OVERVIEW_FIGURES[1]["puml"] = f"""@startdot
digraph G {{
  graph [
    label="图 3-2 用户工作台模块用例图",
    labelloc=t,
    rankdir=LR,
    splines=ortho,
    nodesep=0.45,
    ranksep=1.0,
    pad=0.25,
    fontname="Microsoft YaHei",
    fontsize=28
  ];
  node [fontname="Microsoft YaHei", fontsize=20, color="#334155", penwidth=1.6];
  edge [fontname="Microsoft YaHei", fontsize=18, color="#475569", penwidth=1.8, arrowsize=0.8];

  User [shape=plain, {actor_user}];
  subgraph cluster_system {{
    label="用户工作台模块";
    color="#334155";
    penwidth=1.8;
    fontsize=24;
    style="rounded";
    W1 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="进入与恢复工作台\\n初始化数据/恢复历史/选择场景"];
    W2 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="维护个人资料\\n基本信息/证件照"];
    W3 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="维护教育信息\\n教育背景/核心课程"];
    W4 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="维护经历成果\\n实践经历/项目科研/奖项证书"];
    W5 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="维护能力标签\\n技能能力/熟练度"];
    W6 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="控制工作台界面\\n编辑模块/主题/宽度/AI 面板"];
    W7 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="本地保存与刷新恢复\\nresume-data/page-theme/template"];
    {{ rank=same; W1; W2; W3; W4; W5; W6; W7; }}
    W1 -> W2 [style=invis, weight=50];
    W2 -> W3 [style=invis, weight=50];
    W3 -> W4 [style=invis, weight=50];
    W4 -> W5 [style=invis, weight=50];
    W5 -> W6 [style=invis, weight=50];
    W6 -> W7 [style=invis, weight=50];
  }}
  {{ rank=source; User; }}
  User -> W1 [dir=none];
  User -> W2 [dir=none];
  User -> W3 [dir=none];
  User -> W4 [dir=none];
  User -> W5 [dir=none];
  User -> W6 [dir=none];
  User -> W7 [dir=none];
  W1 -> W7 [style=dashed, arrowhead=open, label="<<include>>"];
  W2 -> W7 [style=dashed, arrowhead=open, label="<<include>>"];
  W3 -> W7 [style=dashed, arrowhead=open, label="<<include>>"];
  W4 -> W7 [style=dashed, arrowhead=open, label="<<include>>"];
  W5 -> W7 [style=dashed, arrowhead=open, label="<<include>>"];
  W6 -> W7 [style=dashed, arrowhead=open, label="<<include>>"];
}}
@enddot
"""
    OVERVIEW_FIGURES[2]["puml"] = f"""@startdot
digraph G {{
  graph [
    label="图 3-3 AI 辅助模块用例图",
    labelloc=t,
    rankdir=LR,
    splines=ortho,
    nodesep=0.45,
    ranksep=1.0,
    pad=0.25,
    fontname="Microsoft YaHei",
    fontsize=28
  ];
  node [fontname="Microsoft YaHei", fontsize=20, color="#334155", penwidth=1.6];
  edge [fontname="Microsoft YaHei", fontsize=18, color="#475569", penwidth=1.8, arrowsize=0.8];

  User [shape=plain, {actor_user}];
  LLM [shape=plain, {actor_llm}];
  subgraph cluster_system {{
    label="AI 辅助模块";
    color="#334155";
    penwidth=1.8;
    fontsize=24;
    style="rounded";
    A1 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="选择优化字段\\n简介/定位/经历/项目"];
    A2 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="AI 文本优化\\n动作/方法/结果/量化成果"];
    A3 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="生成个人简介\\n50-120 字专业总结"];
    A4 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="分析目标要求\\n岗位 JD/院校项目/导师方向"];
    A5 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="匹配材料与目标\\n分数/关键词/缺口/建议"];
    A6 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="生成发展建议\\n短期目标/中期目标/技能路径"];
    A7 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="材料完整度评分\\n本地规则/场景建议"];
    A8 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="AI 调用失败兜底\\n中文错误提示/本地结果"];
    {{ rank=same; A1; A2; A3; A4; A5; A6; A7; A8; }}
    A1 -> A2 [style=invis, weight=50];
    A2 -> A3 [style=invis, weight=50];
    A3 -> A4 [style=invis, weight=50];
    A4 -> A5 [style=invis, weight=50];
    A5 -> A6 [style=invis, weight=50];
    A6 -> A7 [style=invis, weight=50];
    A7 -> A8 [style=invis, weight=50];
  }}
  {{ rank=source; User; }}
  {{ rank=sink; LLM; }}
  User -> A1 [dir=none];
  User -> A2 [dir=none];
  User -> A3 [dir=none];
  User -> A4 [dir=none];
  User -> A5 [dir=none];
  User -> A6 [dir=none];
  User -> A7 [dir=none];
  A2 -> LLM [dir=none];
  A3 -> LLM [dir=none];
  A4 -> LLM [dir=none];
  A5 -> LLM [dir=none];
  A6 -> LLM [dir=none];
  A2 -> A8 [style=dashed, arrowhead=open, label="<<extend>>"];
  A3 -> A8 [style=dashed, arrowhead=open, label="<<extend>>"];
  A4 -> A8 [style=dashed, arrowhead=open, label="<<extend>>"];
  A5 -> A8 [style=dashed, arrowhead=open, label="<<extend>>"];
  A6 -> A8 [style=dashed, arrowhead=open, label="<<extend>>"];
}}
@enddot
"""
    OVERVIEW_FIGURES[3]["puml"] = f"""@startdot
digraph G {{
  graph [
    label="图 3-4 成绩单识别模块用例图",
    labelloc=t,
    rankdir=LR,
    splines=ortho,
    nodesep=0.45,
    ranksep=1.0,
    pad=0.25,
    fontname="Microsoft YaHei",
    fontsize=28
  ];
  node [fontname="Microsoft YaHei", fontsize=20, color="#334155", penwidth=1.6];
  edge [fontname="Microsoft YaHei", fontsize=18, color="#475569", penwidth=1.8, arrowsize=0.8];

  User [shape=plain, {actor_user}];
  OCR [shape=plain, {actor_ocr}];
  subgraph cluster_system {{
    label="成绩单识别模块";
    color="#334155";
    penwidth=1.8;
    fontsize=24;
    style="rounded";
    T1 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="选择成绩单文件\\nPDF/docx/图片格式校验"];
    T2 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="Word 成绩单解析\\ndocument.xml/表格/段落"];
    T3 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="PDF 转图与云端识别\\nPyMuPDF/表格 OCR"];
    T4 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="图片成绩单识别\\nEXIF 修正/压缩/OCR"];
    T5 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="本地文本抽取兜底\\nPDF.js/mammoth/Tesseract"];
    T6 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="候选课程排序筛选\\n成绩/学分/课程名/勾选"];
    T7 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="课程导入教育背景\\n去重/追加 coreCourses"];
    {{ rank=same; T1; T2; T3; T4; T5; T6; T7; }}
    T1 -> T2 [style=invis, weight=50];
    T2 -> T3 [style=invis, weight=50];
    T3 -> T4 [style=invis, weight=50];
    T4 -> T5 [style=invis, weight=50];
    T5 -> T6 [style=invis, weight=50];
    T6 -> T7 [style=invis, weight=50];
  }}
  {{ rank=source; User; }}
  {{ rank=sink; OCR; }}
  User -> T1 [dir=none];
  User -> T6 [dir=none];
  User -> T7 [dir=none];
  T3 -> OCR [dir=none];
  T4 -> OCR [dir=none];
  T1 -> T2 [style=dashed, arrowhead=open, label="<<extend>>"];
  T1 -> T3 [style=dashed, arrowhead=open, label="<<extend>>"];
  T1 -> T4 [style=dashed, arrowhead=open, label="<<extend>>"];
  T2 -> T5 [style=dashed, arrowhead=open, label="<<extend>>"];
  T3 -> T5 [style=dashed, arrowhead=open, label="<<extend>>"];
  T4 -> T5 [style=dashed, arrowhead=open, label="<<extend>>"];
  T2 -> T6 [style=dashed, arrowhead=open, label="<<include>>"];
  T3 -> T6 [style=dashed, arrowhead=open, label="<<include>>"];
  T4 -> T6 [style=dashed, arrowhead=open, label="<<include>>"];
  T5 -> T6 [style=dashed, arrowhead=open, label="<<include>>"];
  T6 -> T7 [style=dashed, arrowhead=open, label="<<include>>"];
}}
@enddot
"""
    OVERVIEW_FIGURES[4]["puml"] = f"""@startdot
digraph G {{
  graph [
    label="图 3-5 预览导出与系统支撑模块用例图",
    labelloc=t,
    rankdir=LR,
    splines=ortho,
    nodesep=0.45,
    ranksep=1.0,
    pad=0.25,
    fontname="Microsoft YaHei",
    fontsize=28
  ];
  node [fontname="Microsoft YaHei", fontsize=20, color="#334155", penwidth=1.6];
  edge [fontname="Microsoft YaHei", fontsize=18, color="#475569", penwidth=1.8, arrowsize=0.8];

  User [shape=plain, {actor_user}];
  Browser [shape=plain, {actor_browser}];
  Ext [shape=plain, {actor_ext}];
  subgraph cluster_system {{
    label="预览导出与系统支撑模块";
    color="#334155";
    penwidth=1.8;
    fontsize=24;
    style="rounded";
    S1 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="模板选择\\n多模板切换/模板缓存"];
    S2 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="实时预览渲染\\n字段兼容/A4 预览"];
    S3 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="PDF 导出\\nhtml2canvas/jsPDF/下载"];
    S4 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="后端接口路由\\noptimize/jd/career/summary/transcript"];
    S5 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="环境变量配置\\.env/VITE_API_BASE/密钥检查"];
    S6 [shape=ellipse, style=filled, fillcolor="#F8FAFC", label="错误提示与服务降级\\n中文提示/保留本地数据"];
    {{ rank=same; S1; S2; S3; S4; S5; S6; }}
    S1 -> S2 [style=invis, weight=50];
    S2 -> S3 [style=invis, weight=50];
    S3 -> S4 [style=invis, weight=50];
    S4 -> S5 [style=invis, weight=50];
    S5 -> S6 [style=invis, weight=50];
  }}
  {{ rank=source; User; }}
  {{ rank=sink; Browser; Ext; }}
  User -> S1 [dir=none];
  User -> S2 [dir=none];
  User -> S3 [dir=none];
  S3 -> Browser [dir=none];
  S4 -> Ext [dir=none];
  S1 -> S2 [style=dashed, arrowhead=open, label="<<include>>"];
  S3 -> S2 [style=dashed, arrowhead=open, label="<<include>>"];
  S4 -> S5 [style=dashed, arrowhead=open, label="<<include>>"];
  S4 -> S6 [style=dashed, arrowhead=open, label="<<extend>>"];
}}
@enddot
"""


def quote(text: str) -> str:
    return text.replace('"', '\\"')


def puml_for_item(number: int, item: dict[str, object]) -> tuple[str, str]:
    caption = f"图 3-{number} {item['title']}时序图"
    participant_defs = participants_for(str(item["kind"]))
    aliases = {alias for _, _, alias in participant_defs}
    actors = {alias for typ, _, alias in participant_defs if typ == "actor"}
    lines = [
        "@startuml",
        "' Generated by scripts/generate_chapter3_plantuml_detailed.py",
        "hide footbox",
        "skinparam dpi 180",
        "skinparam shadowing false",
        "skinparam defaultFontName Microsoft YaHei",
        "skinparam defaultFontSize 20",
        "skinparam titleFontSize 24",
        "skinparam sequenceMessageAlign center",
        "skinparam sequence {",
        "  ArrowThickness 2",
        "  LifeLineBorderColor #64748B",
        "  LifeLineBackgroundColor #F8FAFC",
        "  ParticipantBorderColor #64748B",
        "  ParticipantBackgroundColor #EEF2FF",
        "  ParticipantFontSize 20",
        "  ActorFontSize 20",
        "  ActivateBorderColor #1E293B",
        "  ActivateBackgroundColor #DBEAFE",
        "  BoxPadding 18",
        "}",
        f"title {caption}",
    ]
    for typ, label, alias in participant_defs:
        lines.append(f'{typ} "{quote(label)}" as {alias}')
    lines.append("autonumber")

    active: set[str] = set()
    active_order: list[str] = []

    def activate(alias: str) -> None:
        if alias in actors or alias not in aliases or alias in active:
            return
        lines.append(f"activate {alias}")
        active.add(alias)
        active_order.append(alias)

    def deactivate(alias: str) -> None:
        if alias in active:
            lines.append(f"deactivate {alias}")
            active.remove(alias)

    for start, end, message in item["steps"]:  # type: ignore[index]
        if start == end:
            lines.append(f'{start} -> {end}: {quote(message)}')
            if start not in actors:
                lines.append(f"activate {start}")
                lines.append(f"deactivate {start}")
        else:
            lines.append(f'{start} -> {end}: {quote(message)}')
            if end not in actors:
                activate(end)
            if start in active and start not in actors:
                deactivate(start)
    first_alias = participant_defs[0][2]
    second_alias = participant_defs[1][2]
    lines.extend(
        [
            "alt 正常完成",
            f'{second_alias} --> {first_alias}: {quote(str(item["ok"]))}',
            "else 异常或校验失败",
            f'{second_alias} --> {first_alias}: {quote(str(item["fail"]))}',
            "end",
        ]
    )
    for alias in reversed(active_order):
        deactivate(alias)
    lines.extend(["@enduml", ""])
    return "\n".join(lines), caption


def write_overview_puml_files() -> list[tuple[str, str, Path]]:
    figures: list[tuple[str, str, Path]] = []
    for figure in OVERVIEW_FIGURES:
        stem = str(figure["stem"])
        caption = str(figure["caption"])
        puml_path = PUML_DIR / f"{stem}.puml"
        png_path = DIAGRAM_DIR / f"{stem}.png"
        puml_path.write_text(str(figure["puml"]), encoding="utf-8")
        figure["figure"] = (caption, png_path)
        figures.append(("overview", caption, png_path))
    return figures


def write_puml_files() -> list[tuple[str, str, Path]]:
    figures: list[tuple[str, str, Path]] = []
    figure_no = len(OVERVIEW_FIGURES) + 1
    for module in MODULES:
        for item in module["items"]:  # type: ignore[index]
            puml, caption = puml_for_item(figure_no, item)
            stem = f"figure_3_{figure_no:02d}_{str(item['no']).replace('.', '_')}"
            path = PUML_DIR / f"{stem}.puml"
            path.write_text(puml, encoding="utf-8")
            figures.append((str(item["no"]), caption, DIAGRAM_DIR / f"{stem}.png"))
            item["figure"] = (caption, DIAGRAM_DIR / f"{stem}.png")  # type: ignore[index]
            figure_no += 1
    return figures


def render_plantuml() -> None:
    if not PLANTUML_JAR.exists():
        raise FileNotFoundError(f"PlantUML jar not found: {PLANTUML_JAR}")
    cmd = [
        "java",
        "-DPLANTUML_LIMIT_SIZE=16384",
        "-jar",
        str(PLANTUML_JAR),
        "-charset",
        "UTF-8",
        "-tpng",
        "-o",
        str(DIAGRAM_DIR),
        str(PUML_DIR / "*.puml"),
    ]
    subprocess.run(cmd, check=True, cwd=str(ROOT))


def set_run_font(run, size: float = 12, bold: bool = False, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_style_font(style, chinese: str, size: float) -> None:
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    style.font.size = Pt(size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc(doc: Document) -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("目 录")
    set_run_font(run, 16, True, "黑体")

    toc_paragraph = doc.add_paragraph()
    toc_run = toc_paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "2-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录域：打开 Word 后右键目录并选择“更新域”即可生成页码。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    toc_run._r.append(fld_begin)
    toc_run._r.append(instr)
    toc_run._r.append(fld_sep)
    toc_run._r.append(placeholder)
    toc_run._r.append(fld_end)

    hint = doc.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hint_run = hint.add_run("以下为章节清单，正式页码由上方目录域自动生成。")
    set_run_font(hint_run, 10.5, False, "宋体")

    for module in MODULES:
        p = doc.add_paragraph()
        set_run_font(p.add_run(str(module["title"])), 11, True, "宋体")
        for item in module["items"]:  # type: ignore[index]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(22)
            set_run_font(p.add_run(f"{item['no']} {item['title']}"), 10.5, False, "宋体")
    doc.add_page_break()


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 12, False, "宋体")


def add_figure(doc: Document, caption: str, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    with Image.open(path) as image:
        width_px, height_px = image.size
    height_at_max_width = MAX_FIGURE_WIDTH_CM * height_px / width_px
    if height_at_max_width > MAX_FIGURE_HEIGHT_CM:
        p.add_run().add_picture(str(path), height=Cm(MAX_FIGURE_HEIGHT_CM))
    else:
        p.add_run().add_picture(str(path), width=Cm(MAX_FIGURE_WIDTH_CM))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run_font(run, 10.5, False, "宋体")


def write_docx() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_page_number(section.footer.paragraphs[0])

    set_style_font(doc.styles["Normal"], "宋体", 12)
    set_style_font(doc.styles["Heading 1"], "黑体", 16)
    set_style_font(doc.styles["Heading 2"], "黑体", 14)
    set_style_font(doc.styles["Heading 3"], "黑体", 12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("3. 需求分析——功能建模")
    set_run_font(run, 18, True, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("项目名称：AI 简历与申请材料工作台"), 12, False, "宋体")

    add_toc(doc)

    add_paragraph(
        doc,
        "本章依据 AI 简历与申请材料工作台的当前代码实现进行功能建模。系统面向材料编写者，支持求职、考研、保研、升学四类材料场景，核心能力包括结构化编辑、实时预览、AI 辅助优化、成绩单识别、材料评分、PDF 导出和异常降级。为与课程报告规范保持一致，本章按大模块—三级功能点—时序图的方式展开，每个三级功能点均描述触发条件、系统处理、正常结果和异常分支。",
    )
    add_paragraph(
        doc,
        "当前系统未实现账号登录、管理员后台或服务器端数据持久化，因此本文不虚构用户认证、权限审核等功能，而是将已有前端工作台、AI 服务、OCR 解析和导出支撑能力拆分为 4 个一级模块、36 个三级功能点。",
    )
    add_paragraph(
        doc,
        "参照软件工程需求分析报告的常见写法，本章先给出系统级与模块级用例图，明确参与者、系统边界和主要用例；再给出总体流程图与关键业务流程图，说明业务处理顺序；最后逐项给出带活跃期的 UML 时序图，细化每个功能点内部对象之间的交互。",
    )

    for figure in OVERVIEW_FIGURES:
        add_paragraph(doc, str(figure["desc"]))
        caption, figure_path = figure["figure"]  # type: ignore[index]
        add_figure(doc, caption, Path(figure_path))

    for module in MODULES:
        doc.add_heading(str(module["title"]), level=2)
        add_paragraph(doc, str(module["intro"]))
        for item in module["items"]:  # type: ignore[index]
            doc.add_heading(f"{item['no']} {item['title']}", level=3)
            add_paragraph(doc, str(item["desc"]))
            caption, figure_path = item["figure"]  # type: ignore[index]
            add_figure(doc, caption, Path(figure_path))

    output = DOCS / "第三章_需求分析_功能建模.docx"
    doc.save(output)
    return output


def write_markdown() -> Path:
    lines = [
        "# 3. 需求分析——功能建模",
        "",
        "项目名称：AI 简历与申请材料工作台",
        "",
        "本章按 4 个一级模块、36 个三级功能点展开，先给出系统级与模块级用例图、关键业务流程图，再为每个功能点配置带活跃期的 PlantUML 时序图。",
        "",
    ]
    for figure in OVERVIEW_FIGURES:
        lines.append(str(figure["desc"]))
        lines.append("")
        caption, figure_path = figure["figure"]  # type: ignore[index]
        rel = Path(figure_path).relative_to(DOCS).as_posix()
        lines.append(f"![{caption}]({rel})")
        lines.append("")
        lines.append(f"<div align=\"center\">{caption}</div>")
        lines.append("")
    for module in MODULES:
        lines.append(f"## {module['title']}")
        lines.append("")
        lines.append(str(module["intro"]))
        lines.append("")
        for item in module["items"]:  # type: ignore[index]
            lines.append(f"### {item['no']} {item['title']}")
            lines.append("")
            lines.append(str(item["desc"]))
            lines.append("")
            caption, figure_path = item["figure"]  # type: ignore[index]
            rel = Path(figure_path).relative_to(DOCS).as_posix()
            lines.append(f"![{caption}]({rel})")
            lines.append("")
            lines.append(f"<div align=\"center\">{caption}</div>")
            lines.append("")
    output = DOCS / "第三章_需求分析_功能建模.md"
    output.write_text("\n".join(lines), encoding="utf-8")
    return output


def main() -> None:
    ensure_dirs()
    apply_report_use_case_templates()
    apply_ordered_use_case_templates()
    apply_report_flow_templates()
    figures = write_overview_puml_files()
    figures.extend(write_puml_files())
    render_plantuml()
    missing = [path for _, _, path in figures if not path.exists()]
    if missing:
        raise RuntimeError(f"missing rendered diagrams: {missing[:3]}")
    docx = write_docx()
    md = write_markdown()
    print(f"PlantUML diagrams: {len(figures)}")
    print(f"Word report: {docx}")
    print(f"Markdown report: {md}")
    print(f"PUML sources: {PUML_DIR}")


if __name__ == "__main__":
    main()
