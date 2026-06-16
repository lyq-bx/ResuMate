from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\BiographicalNotes\exp4_work")
ROOT.mkdir(parents=True, exist_ok=True)


def copy_assets():
    source = Path(r"D:\自然语言处理\exp4")
    by_size = {
        251586: "bert_base.png",
        282370: "bert_mrc.png",
        293317: "bert_task_compare.png",
        295148: "bert_classification.png",
        320696: "bert_ner.png",
        1573394: "cover_visual.png",
    }
    for image in source.glob("*.png"):
        name = by_size.get(image.stat().st_size)
        if name:
            (ROOT / name).write_bytes(image.read_bytes())


def generate_arch_image():
    out = ROOT / "model_arch_comparison.png"
    font_path = r"C:\Windows\Fonts\msyh.ttc"
    font_manager.fontManager.addfont(font_path)
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei"]
    plt.rcParams["axes.unicode_minus"] = False

    fig = plt.figure(figsize=(15, 8.5), dpi=220)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 8.5)
    ax.axis("off")
    fig.patch.set_facecolor("#f7f8fb")

    ax.text(
        7.5,
        8.05,
        "预训练语言模型架构对比",
        ha="center",
        va="center",
        fontsize=24,
        fontweight="bold",
        color="#1f2937",
    )
    ax.text(
        7.5,
        7.65,
        "BERT 侧重理解，GPT-2 侧重自回归生成，T5 统一为文本到文本的编码器-解码器框架",
        ha="center",
        va="center",
        fontsize=12.5,
        color="#4b5563",
    )

    colors = {
        "bert": ("#e8f3ff", "#1d4ed8"),
        "gpt": ("#fff7e6", "#b45309"),
        "t5": ("#eafaf1", "#047857"),
        "muted": "#6b7280",
    }

    def rounded(x, y, w, h, fc, ec, text="", fs=11, weight="normal", color="#111827"):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.03,rounding_size=0.08",
            linewidth=1.4,
            facecolor=fc,
            edgecolor=ec,
        )
        ax.add_patch(patch)
        if text:
            ax.text(
                x + w / 2,
                y + h / 2,
                text,
                ha="center",
                va="center",
                fontsize=fs,
                fontweight=weight,
                color=color,
                wrap=True,
            )
        return patch

    def arrow(x1, y1, x2, y2, color="#6b7280"):
        ax.add_patch(
            FancyArrowPatch(
                (x1, y1),
                (x2, y2),
                arrowstyle="-|>",
                mutation_scale=12,
                linewidth=1.3,
                color=color,
            )
        )

    cols = [0.65, 5.2, 9.75]
    labels = [
        ("BERT", "Encoder-only", "双向自注意力", "bert"),
        ("GPT-2", "Decoder-only", "因果自注意力", "gpt"),
        ("T5", "Encoder-Decoder", "文本到文本", "t5"),
    ]
    for x, (name, typ, sub, key) in zip(cols, labels):
        _, ec = colors[key]
        rounded(x, 0.65, 4.0, 6.65, "#ffffff", "#d1d5db")
        ax.add_patch(Rectangle((x, 6.78), 4.0, 0.08, color=ec, lw=0))
        ax.text(x + 0.28, 6.45, name, fontsize=21, fontweight="bold", color=ec, ha="left")
        ax.text(x + 0.28, 6.12, typ + " | " + sub, fontsize=11.5, color=colors["muted"], ha="left")

    x = cols[0]
    fc, ec = colors["bert"]
    rounded(x + 0.35, 5.35, 3.3, 0.55, fc, ec, "[CLS] 文本A [SEP] 文本B [SEP]", 9.7)
    arrow(x + 2.0, 5.35, x + 2.0, 4.98, ec)
    for i in range(4):
        rounded(x + 0.85, 4.45 - i * 0.48, 2.3, 0.34, "#eff6ff", ec, "Transformer Encoder", 9)
    ax.text(x + 3.22, 3.58, "×12", fontsize=13, fontweight="bold", color=ec)
    arrow(x + 2.0, 2.72, x + 2.0, 2.35, ec)
    rounded(x + 0.35, 1.65, 1.55, 0.55, "#f8fbff", ec, "[CLS] 分类", 9.5)
    rounded(x + 2.1, 1.65, 1.55, 0.55, "#f8fbff", ec, "Token/Span 输出", 9.5)
    ax.text(x + 0.35, 1.15, "适用：情感分类、文本蕴含、阅读理解、NER", fontsize=9.7, color="#374151", ha="left")
    ax.text(x + 0.35, 0.88, "训练目标：MLM/NSP 预训练 + 下游任务监督微调", fontsize=9.3, color=colors["muted"], ha="left")

    x = cols[1]
    fc, ec = colors["gpt"]
    rounded(x + 0.35, 5.35, 3.3, 0.55, fc, ec, "上文 tokens", 10.5)
    arrow(x + 2.0, 5.35, x + 2.0, 4.98, ec)
    for i in range(4):
        rounded(x + 0.85, 4.45 - i * 0.48, 2.3, 0.34, "#fffaf0", ec, "Masked Decoder Block", 9)
    ax.text(x + 3.22, 3.58, "×N", fontsize=13, fontweight="bold", color=ec)
    arrow(x + 2.0, 2.72, x + 2.0, 2.35, ec)
    rounded(x + 0.85, 1.65, 2.3, 0.55, "#fffbeb", ec, "预测下一个 token", 10)
    ax.text(x + 0.35, 1.15, "适用：续写、开放式文本生成、语言建模", fontsize=9.7, color="#374151", ha="left")
    ax.text(x + 0.35, 0.88, "训练目标：Causal LM，只看左侧上下文", fontsize=9.3, color=colors["muted"], ha="left")

    x = cols[2]
    fc, ec = colors["t5"]
    rounded(x + 0.35, 5.45, 3.3, 0.5, fc, ec, "任务前缀 + 输入文本", 10)
    arrow(x + 2.0, 5.45, x + 2.0, 5.08, ec)
    rounded(x + 0.55, 4.48, 2.9, 0.42, "#f0fdf4", ec, "Encoder：理解源序列", 9.5)
    rounded(x + 0.55, 3.82, 2.9, 0.42, "#f0fdf4", ec, "Cross Attention：对齐源/目标", 9.5)
    rounded(x + 0.55, 3.16, 2.9, 0.42, "#f0fdf4", ec, "Decoder：逐步生成目标", 9.5)
    arrow(x + 2.0, 3.16, x + 2.0, 2.35, ec)
    rounded(x + 0.85, 1.65, 2.3, 0.55, "#ecfdf5", ec, "输出文本序列", 10)
    ax.text(x + 0.35, 1.15, "适用：翻译、摘要、问答、格式转换", fontsize=9.7, color="#374151", ha="left")
    ax.text(x + 0.35, 0.88, "训练目标：Seq2Seq，输入输出都表示为文本", fontsize=9.3, color=colors["muted"], ha="left")

    rounded(0.65, 0.18, 13.1, 0.34, "#ffffff", "#d1d5db")
    ax.text(
        1.0,
        0.35,
        "对比结论：BERT 的优势是双向语义表征；GPT-2 的优势是自然流畅地续写；T5 兼具理解与生成，适合需要输入到输出映射的任务。",
        fontsize=10,
        color="#374151",
        va="center",
    )

    fig.savefig(out, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def set_text_font(run, size=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_text_font(r, 10, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_text_font(r, 16 if level == 1 else 13, True)
    r.font.color.rgb = RGBColor(31, 78, 121) if level == 1 else RGBColor(46, 116, 181)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_text_font(r, 10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(item)
        set_text_font(r, 10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(item)
        set_text_font(r, 10.5)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    for line in code.splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9)


def add_picture(doc, filename, caption, width_cm=14.0):
    path = ROOT / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_text_font(r, 9.5)
    r.italic = True


def build_report():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(2.1)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("自然语言处理\n实验报告四")
    set_text_font(r, 26, True)

    doc.add_paragraph()
    doc.add_paragraph()
    info = [
        ("名称", "预训练语言模型"),
        ("专业", "智能科学与技术"),
        ("年级", "大三"),
        ("学号", "20231060024"),
        ("姓名", "吴宣萱"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        set_cell_text(table.cell(i, 0), f"{k}：", True)
        set_cell_text(table.cell(i, 1), v)
        table.cell(i, 0).width = Cm(5)
        table.cell(i, 1).width = Cm(7)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("06/2026")
    set_text_font(r, 14)
    doc.add_page_break()

    add_heading(doc, "一、实验目的及要求")
    add_para(
        doc,
        "本次实验的主要目的是熟练掌握 Hugging Face Transformers 库的使用方法，理解预训练语言模型在自然语言理解和自然语言生成任务中的微调流程，并通过 BERT、GPT-2、T5 三类模型完成不同下游任务。",
    )
    add_bullets(
        doc,
        [
            "掌握 BERT 的输入表示、Transformer 编码器结构、预训练任务与微调方法。",
            "完成基于 BERT 的单句分类、句对分类、抽取式阅读理解和命名实体识别任务。",
            "完成基于 GPT-2 的条件文本生成任务，理解自回归语言模型的训练目标。",
            "完成基于 T5/mT5 的机器翻译任务，理解编码器-解码器模型和文本到文本范式。",
            "结合训练日志、评价指标和输出样例，对不同任务的模型结构、训练效果和误差来源进行分析。",
        ],
    )

    add_heading(doc, "二、实验内容")
    add_table(
        doc,
        ["编号", "模型任务", "数据集", "主要评价指标"],
        [
            ("实验一", "BERT 单句文本分类", "SST-2 情感分析", "Accuracy"),
            ("实验二", "BERT 句对文本分类", "GLUE RTE 文本蕴含", "Accuracy"),
            ("实验三", "BERT 抽取式阅读理解", "SQuAD 问答抽取", "EM / F1 / Eval Loss"),
            ("实验四", "BERT 序列标注", "CoNLL-2003 NER", "Precision / Recall / F1 / Accuracy"),
            ("实验五", "GPT-2 文本生成", "WikiText-2 语言建模", "Perplexity"),
            ("实验六", "T5/mT5 机器翻译", "IWSLT2017 中英翻译", "BLEU / Loss / 生成样例"),
        ],
        [2.0, 4.0, 4.2, 4.0],
    )

    add_heading(doc, "三、模型架构对比")
    add_picture(doc, "model_arch_comparison.png", "图1  BERT、GPT-2 与 T5 的总体架构对比", 15.5)
    add_table(
        doc,
        ["模型", "结构类型", "注意力方向", "典型输入/输出", "适合任务", "本实验对应任务"],
        [
            ("BERT", "Encoder-only", "双向自注意力", "输入文本 -> 语义表示/标签", "分类、匹配、抽取、序列标注", "SSC、SPC、MRC、NER"),
            ("GPT-2", "Decoder-only", "从左到右的因果注意力", "上文 -> 下一个 token/后续文本", "开放式生成、语言建模", "WikiText 条件文本生成"),
            ("T5/mT5", "Encoder-Decoder", "编码端双向，解码端因果，带交叉注意力", "带任务前缀的输入文本 -> 输出文本", "翻译、摘要、问答、改写", "IWSLT 机器翻译"),
        ],
        [2.0, 2.7, 3.0, 3.8, 3.1, 3.0],
    )
    add_para(
        doc,
        "对比来看，BERT 的优势在于利用双向上下文形成稳定的语义表示，因此适合理解类任务；GPT-2 使用因果掩码逐词预测，天然适合文本续写和语言建模；T5/mT5 使用编码器-解码器结构，将不同任务统一为“输入文本到输出文本”的形式，在翻译、摘要等生成式任务中更灵活。",
    )
    add_picture(doc, "bert_task_compare.png", "图2  BERT 在不同理解任务中的输出头对比", 14.5)

    add_heading(doc, "四、实验结果、代码解析与分析")

    add_heading(doc, "4.1 实验一：BERT 实现情感分析的单句文本分类任务", 2)
    add_para(doc, "本实验使用 SST-2 数据集完成二分类情感分析。输入为单个句子，输出为 positive/negative 两类标签。模型主体为 bert-base-cased，任务输出层为 BertForSequenceClassification 中的线性分类头。")
    add_code(
        doc,
        """dataset = load_dataset('nyu-mll/glue', 'sst2')
tokenizer = BertTokenizerFast.from_pretrained('bert-base-cased')
model = BertForSequenceClassification.from_pretrained(
    'bert-base-cased', num_labels=2, return_dict=True
)
def tokenize(examples):
    return tokenizer(examples['sentence'], truncation=True, padding='max_length')""",
    )
    add_bullets(
        doc,
        [
            "代码首先加载 SST-2 数据集、BERT 分词器和二分类模型。SST-2 是单句情感分类任务，因此只需要对 sentence 字段进行编码。",
            "tokenizer 会生成 input_ids、token_type_ids 和 attention_mask，其中 attention_mask 用于屏蔽 padding 部分，避免填充 token 影响模型计算。",
            "BertForSequenceClassification 取 [CLS] 位置的 pooled output，经 dropout 和线性层得到两个类别的 logits。",
            "compute_metrics 使用 np.argmax 取最大概率类别，再调用 GLUE/SST-2 的 accuracy 指标计算验证集准确率。",
            "由于训练时出现过内存不足，代码将 batch size 调小，并使用 gradient_accumulation_steps=2 与 fp16=True，在显存/内存较紧张时保持近似的有效批量大小。",
        ],
    )
    add_para(
        doc,
        "结果分析：从训练日志和损失曲线看，模型在前期能够较快下降，说明预训练 BERT 已经具备较强的通用语义表示能力，微调阶段主要是在学习情感分类边界。SST-2 句子较短、标签数量少，因此相比阅读理解和翻译任务收敛更快。若验证准确率后期提升变慢或波动，主要原因可能是部分句子情感表达含蓄、存在否定转折，或者训练轮数较少导致分类头尚未完全稳定。整体上，单句分类任务最能体现 BERT 的 [CLS] 表征在句级语义任务中的迁移能力。",
    )
    add_picture(doc, "bert_classification.png", "图3  BERT 单句/句对分类模型结构", 13.5)

    add_heading(doc, "4.2 实验二：BERT 实现文本蕴含的句对文本分类任务", 2)
    add_para(doc, "本实验使用 GLUE RTE 数据集完成文本蕴含判断。输入为 sentence1 和 sentence2 两个句子，输出为 entailment 或 not entailment。")
    add_code(
        doc,
        """dataset = load_dataset('glue', 'rte')
tokenizer = BertTokenizerFast.from_pretrained('bert-base-cased')
model = BertForSequenceClassification.from_pretrained('bert-base-cased', return_dict=True)
def tokenize(examples):
    return tokenizer(examples['sentence1'], examples['sentence2'],
                     truncation=True, padding='max_length')""",
    )
    add_bullets(
        doc,
        [
            "句对任务的输入格式为 [CLS] sentence1 [SEP] sentence2 [SEP]，token_type_ids 用于区分两个句子片段。",
            "与单句分类相比，模型不仅要理解每个句子的语义，还要判断两个句子之间是否存在蕴含关系。",
            "训练参数仍采用较小学习率 2e-5 和 2 个 epoch，符合预训练模型微调时“低学习率、少轮数”的常用策略。",
            "评价函数调用 GLUE/RTE 指标，主要观察 accuracy。",
        ],
    )
    add_para(
        doc,
        "结果分析：RTE 数据规模较小，且推理关系比情感分类更复杂，因此验证集指标通常比 SST-2 更容易波动。BERT 的 segment embedding 和双向注意力可以同时建模两个句子的词语对应关系，例如同义改写、否定关系和上下位关系；但当句子中存在复杂常识推理或长距离依赖时，单纯的二分类头可能仍然不足。实验结果说明，句对分类任务对语义匹配能力要求更高，训练曲线不一定像单句分类那样平滑。",
    )

    add_heading(doc, "4.3 实验三：BERT 实现抽取式阅读理解任务", 2)
    add_para(doc, "本实验使用 SQuAD 数据集完成抽取式问答。模型输入由问题和篇章组成，输出不是类别标签，而是答案在篇章中的起始 token 位置和结束 token 位置。")
    add_code(
        doc,
        """dataset = load_dataset('rajpurkar/squad')
tokenizer = BertTokenizerFast.from_pretrained('bert-base-cased')
model = BertForQuestionAnswering.from_pretrained(
    'bert-base-cased', return_dict=True
)
def prepare_train_features(examples):
    tokenized_examples = tokenizer(
        examples['question'], examples['context'],
        truncation='only_second', max_length=384, stride=128,
        return_overflowing_tokens=True,
        return_offsets_mapping=True,
        padding='max_length'
    )""",
    )
    add_bullets(
        doc,
        [
            "代码将 question 作为第一段、context 作为第二段，并设置 truncation=\"only_second\"，保证问题不会被截断，长篇章只截断 context。",
            "max_length=384 与 stride=128 用于处理长篇章：当 context 超过最大长度时切成多个重叠窗口，避免答案刚好落在截断边界外。",
            "return_offsets_mapping 建立 token 与原始字符位置的映射，用于把人工标注的字符级答案位置转换为 token 级 start_positions 和 end_positions。",
            "BertForQuestionAnswering 在每个 token 上输出 start logits 和 end logits，训练时分别计算起点、终点的交叉熵损失。",
            "本 notebook 中主要使用 eval_loss 观察训练效果；若要得到标准 SQuAD EM/F1，需要在预测后将 start/end logits 后处理为文本答案，再调用 squad metric。",
        ],
    )
    add_para(
        doc,
        "结果分析：阅读理解任务比分类任务更细粒度，因为模型必须在篇章中准确定位答案边界。训练过程中，若训练损失和验证损失同步下降，说明模型正在学习问题与上下文片段之间的对齐关系；若验证损失下降缓慢，常见原因是长篇章被切片后答案窗口匹配复杂，或答案边界存在多个合理表达。SQuAD 的答案通常是原文连续片段，因此 BERT 的双向编码器适合同时利用答案左右两侧上下文，但它不能生成原文中不存在的答案。",
    )
    add_picture(doc, "bert_mrc.png", "图4  BERT 抽取式阅读理解模型结构", 13.5)

    add_heading(doc, "4.4 实验四：BERT 实现命名实体识别的序列标注任务", 2)
    add_para(doc, "本实验使用 CoNLL-2003 数据集完成 NER。与句级分类不同，NER 需要对输入序列中的每个 token 输出一个 BIO 标签，例如 B-PER、I-ORG、B-LOC 等。")
    add_code(
        doc,
        """label2id = {
    'O': 0, 'B-MISC': 1, 'I-MISC': 2, 'B-PER': 3, 'I-PER': 4,
    'B-ORG': 5, 'I-ORG': 6, 'B-LOC': 7, 'I-LOC': 8
}
tokenized_inputs = tokenizer(examples['tokens'], truncation=True,
                             is_split_into_words=True)
word_ids = tokenized_inputs.word_ids(batch_index=i)
# 特殊 token 标为 -100，计算 loss 时自动忽略
model = BertForTokenClassification.from_pretrained(
    'bert-base-cased', num_labels=len(label_list),
    id2label=id2label, label2id=label2id
)""",
    )
    add_bullets(
        doc,
        [
            "代码从 GitHub 镜像下载 CoNLL-2003 的 train/validation/test 文本文件，并定义 9 类 BIO 标签映射。",
            "tokenizer 使用 is_split_into_words=True，表示输入已经按词切分。BERT WordPiece 可能把一个词拆成多个子词，因此需要通过 word_ids 将原始词标签对齐到子词。",
            "特殊符号 [CLS]、[SEP] 和 padding 的标签设置为 -100，PyTorch 交叉熵会忽略这些位置，避免它们参与损失计算。",
            "DataCollatorForTokenClassification 负责动态 padding，并保持 labels 与 input_ids 长度一致。",
            "seqeval 计算实体级 precision、recall、F1 和 accuracy，比单纯 token accuracy 更能反映实体边界和实体类型是否正确。",
        ],
    )
    add_para(
        doc,
        "结果分析：NER 的关键难点在于实体边界识别和实体类型区分。训练后若 F1 明显提升，说明 BERT 的上下文表示能够帮助判断同一个词在不同语境中的实体类别。例如 “Washington” 在不同句子中可能表示地点、组织或人名。常见错误包括多词实体边界截断、MISC 类别样本较少导致召回不足，以及 WordPiece 切分后子词标签不稳定。总体而言，BERT + token classification head 是序列标注任务中较直接有效的微调方式。",
    )
    add_picture(doc, "bert_ner.png", "图5  BERT 命名实体识别模型结构", 13.5)

    add_heading(doc, "4.5 实验五：GPT-2 实现面向维基百科的条件文本生成任务", 2)
    add_para(doc, "本实验使用 WikiText-2 数据集对 GPT-2 类自回归语言模型进行微调。代码中分词器使用 gpt2，模型使用 distilgpt2；两者共享 GPT-2 系列词表，因此可以配合使用。")
    add_code(
        doc,
        """wikitext_data = load_dataset('wikitext', 'wikitext-2-v1')
tokenizer = AutoTokenizer.from_pretrained('gpt2')
model = AutoModelForCausalLM.from_pretrained('distilgpt2')

def group_texts(examples):
    concatenated = {k: sum(examples[k], []) for k in examples.keys()}
    result = {k: [t[i:i+128] for i in range(0, total_length, 128)]
              for k, t in concatenated.items()}
    result['labels'] = result['input_ids'].copy()""",
    )
    add_bullets(
        doc,
        [
            "preprocess_function 将原始文本转为 token 序列，group_texts 将连续文本切成 block_size=128 的训练块。",
            "labels 直接复制 input_ids，表示模型在每个位置预测下一个 token，这是因果语言模型的核心训练目标。",
            "DataCollatorForLanguageModeling 设置 mlm=False，说明不使用 BERT 式遮蔽预测，而是使用 GPT 式自回归预测。",
            "评价阶段使用 perplexity = exp(eval_loss)。困惑度越低，表示模型对测试文本的平均预测不确定性越小。",
        ],
    )
    add_para(
        doc,
        "结果分析：GPT-2 的训练目标与 BERT 明显不同。BERT 在预训练时双向看上下文并预测被遮蔽词，而 GPT-2 只能看当前位置左侧文本，逐 token 生成后续内容。训练后如果 eval_loss 降低、perplexity 下降，说明模型对维基百科风格文本的续写分布更加熟悉。生成质量不仅取决于 loss，还取决于采样策略，例如 top_k、top_p、temperature 等；采样过保守会导致文本重复，采样过随机则可能引入语义漂移。",
    )

    add_heading(doc, "4.6 实验六：T5/mT5 实现机器翻译任务", 2)
    add_para(doc, "本 notebook 实际加载的是 iwslt2017-zh-en 数据集和 google/mt5-small 模型，因此本节按中英机器翻译任务分析。若更换为英法数据集，整体流程仍然一致：输入源语言句子，输出目标语言句子。")
    add_code(
        doc,
        """model_name = 'google/mt5-small'
iwslt_data = load_dataset('iwslt2017', 'iwslt2017-zh-en')
prefix = 'translate Chinese to English: '

def preprocess_function(examples):
    inputs = [prefix + ex['zh'] for ex in examples['translation']]
    targets = [ex['en'] for ex in examples['translation']]
    return tokenizer(inputs, text_target=targets,
                     max_length=128, truncation=True)""",
    )
    add_bullets(
        doc,
        [
            "T5/mT5 采用文本到文本范式，任务通过前缀 translate Chinese to English: 显式告诉模型当前目标是翻译。",
            "tokenizer(inputs, text_target=targets) 同时编码源语言和目标语言，训练时 decoder 根据目标序列计算交叉熵损失。",
            "DataCollatorForSeq2Seq 会自动处理 encoder 输入、decoder 输入和标签 padding，避免手工拼接 decoder_input_ids。",
            "Seq2SeqTrainer 支持生成式评估，compute_metrics 中可使用 sacrebleu 计算 BLEU 分数。当前训练参数中 do_eval=False，主要通过训练损失和翻译样例观察效果。",
            "训练完成后从 checkpoint-10000 加载模型，使用 generate 生成译文，并通过 top_k/top_p 采样控制译文多样性。",
        ],
    )
    add_para(
        doc,
        "结果分析：已记录的训练曲线显示，前 0~5000 步损失从约 13.4 快速下降，说明 mT5 预训练权重能够迅速适配翻译任务；5000~20000 步下降速度放缓，模型开始学习更细致的词汇、语序和句式映射；20000~31500 步后损失稳定在约 2.94~3.00 区间小幅震荡，说明模型基本收敛，继续训练带来的收益有限。机器翻译是序列生成任务，交叉熵损失数值通常高于二分类任务，因此应重点观察下降趋势、验证集 BLEU 和实际译文可读性，而不是只比较 loss 的绝对大小。",
    )

    add_heading(doc, "五、问题回答")
    add_heading(doc, "5.1 简要说明预训练语言模型的三大特点", 2)
    add_numbered(
        doc,
        [
            "大规模自监督预训练：预训练语言模型通常在大规模无标注语料上训练，通过 MLM、Causal LM、Seq2Seq 等目标学习通用语言规律。",
            "上下文相关表示能力：同一个词在不同句子中会得到不同表示，模型能够结合上下文理解词义、句义和篇章关系。",
            "迁移与微调能力强：模型预训练后可通过少量任务数据进行微调，迁移到分类、问答、标注、生成、翻译等多种下游任务。",
        ],
    )
    add_heading(doc, "5.2 简要说明 GPT 和 BERT 模型的异同", 2)
    add_para(doc, "相同点：GPT 和 BERT 都基于 Transformer，都先在大规模语料上预训练，再迁移到下游任务；两者都使用分词器、词向量、位置向量、多层自注意力和前馈网络。")
    add_para(doc, "不同点：BERT 是 Encoder-only 结构，使用双向自注意力，预训练目标主要是 MLM/NSP，适合理解类任务；GPT 是 Decoder-only 结构，使用因果自注意力，只能看到左侧上下文，预训练目标是预测下一个 token，适合文本生成任务。")
    add_heading(doc, "5.3 BERT 输入表示为什么要包含位置向量", 2)
    add_para(doc, "Transformer 的自注意力本身对输入顺序不敏感。如果没有位置向量，模型只能看到一组 token，而难以判断词语先后关系。例如“狗追人”和“人追狗”包含相同词语，但语义完全不同。位置向量为每个 token 注入序列位置信息，使 BERT 能够区分词序、短语结构和长距离依赖。若去掉位置向量，模型在语序敏感任务上的效果会明显下降，尤其会影响阅读理解、句法关系判断和实体边界识别。")
    add_heading(doc, "5.4 MLM、WWM 和 NM 掩码策略的异同", 2)
    add_para(doc, "三种策略的共同点是都用于 BERT 类模型的预训练阶段，通过遮蔽部分输入并要求模型恢复被遮蔽内容，让模型学习上下文语义表示；在下游任务微调阶段，模型结构通常不需要改变，仍然加载预训练权重并接任务输出层。")
    add_table(
        doc,
        ["策略", "预训练阶段差异", "下游精调阶段影响"],
        [
            ("MLM", "随机遮蔽若干 token，可能只遮蔽一个词的一部分 WordPiece 子词。", "流程最标准，适用于通用 BERT 微调。"),
            ("WWM", "Whole Word Masking，若一个词被选中，则该词拆出的所有子词一起遮蔽。", "词级语义学习更完整，中文或子词较多场景常能提升理解任务效果。"),
            ("NM", "通常指 N-gram Masking，连续遮蔽若干 token 或词片段，增强模型对短语级连续片段的建模。", "微调方式与 MLM/WWM 基本相同，但预训练权重可能更擅长短语恢复和片段级语义理解。"),
        ],
        [2.0, 7.0, 6.0],
    )

    add_heading(doc, "六、遇到的问题和解决方案")
    add_bullets(
        doc,
        [
            "数据集加载路径问题：SQuAD 加载时出现 URI 或命名空间问题，解决方法是使用 rajpurkar/squad 这样的完整数据集路径，并更新 datasets 与 huggingface_hub。",
            "运行内存不足：SST-2 训练中出现 OOM/RAM 不足，解决方法是减小 per_device_train_batch_size，使用 gradient_accumulation_steps 保持有效批量，并开启 fp16 混合精度。",
            "NER 数据集脚本依赖问题：CoNLL-2003 直接从 Hugging Face 加载不稳定，改为从 GitHub 下载官方文本文件并本地解析，保证实验可复现。",
            "生成式任务评估开销较大：T5/mT5 翻译训练中暂时关闭完整验证集评估，通过保存 checkpoint、观察训练损失和抽样翻译结果判断模型是否收敛。",
            "不同任务输出形式差异较大：分类任务输出句级标签，阅读理解输出 span 起止位置，NER 输出 token 标签，GPT-2/T5 输出文本序列。实验中需要分别选择对应的模型类、数据整理函数和评价指标。",
        ],
    )

    add_heading(doc, "七、实验总结")
    add_para(
        doc,
        "通过本次实验，可以清晰看到预训练语言模型在不同 NLP 任务中的迁移方式：BERT 通过更换任务输出头即可完成分类、句对匹配、阅读理解和序列标注；GPT-2 通过自回归目标完成开放式文本生成；T5/mT5 将翻译任务统一为文本到文本生成。不同模型的架构差异决定了它们的优势场景，理解类任务更适合双向编码器，连续文本生成更适合因果解码器，而机器翻译等输入输出映射任务更适合编码器-解码器结构。实验也说明，微调不仅需要关注模型本身，还要重视分词对齐、长文本切片、batch size、评价指标和生成策略等工程细节。",
    )
    add_picture(doc, "bert_base.png", "图6  BERT 基础架构示意图", 13.5)

    out = ROOT / "completed_report_utf8.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    copy_assets()
    generate_arch_image()
    report = build_report()
    print(report)
